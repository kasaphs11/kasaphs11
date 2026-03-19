# word_export.py
# Export προγραμμάτων σε Word (.docx)

import calendar
import datetime as dt

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls
from docx.oxml.ns import qn

from logic.constants    import GREEK_MONTHS_GEN, GREEK_MONTH_ABBR, TAB_TITLES
from logic.date_helpers import day_bucket

# -----------------------------
# Export to Word
# -----------------------------
def set_cell_background(cell, color_rgb: tuple[int, int, int]):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), f'{color_rgb[0]:02X}{color_rgb[1]:02X}{color_rgb[2]:02X}')
    cell._element.get_or_add_tcPr().append(shading)



def export_all_schedules_to_word(
    filepath: str,
    export_data: list[dict],
    year: int,
    month: int,
    extra_holidays: set[int],
):
    """
    Export all schedules to a single Word file with page breaks.
    """
    from docx.oxml import OxmlElement as _OE

    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2)
        section.right_margin  = Cm(2)

    def set_font(run, size_pt, bold=False):
        run.font.name = "Cambria"
        run.font.size = Pt(size_pt)
        run.bold = bold

    def set_row_height(row, h_cm=0.5):
        trPr = row._tr.get_or_add_trPr()
        trH = _OE("w:trHeight")
        trH.set(qn("w:val"), str(int(Cm(h_cm).twips)))
        trH.set(qn("w:hRule"), "exact")
        trPr.append(trH)

    def set_cell_width(cell, w_cm):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = _OE("w:tcW")
        tcW.set(qn("w:w"), str(int(Cm(w_cm).twips)))
        tcW.set(qn("w:type"), "dxa")
        existing = tcPr.find(qn("w:tcW"))
        if existing is not None:
            tcPr.remove(existing)
        tcPr.insert(0, tcW)

    def set_cell_valign_center(cell):
        tcPr = cell._tc.get_or_add_tcPr()
        vAlign = _OE("w:vAlign")
        vAlign.set(qn("w:val"), "center")
        tcPr.append(vAlign)

    col_widths = [3, 3, 3, 5.5]
    _, days_in_month = calendar.monthrange(year, month)

    for idx, data in enumerate(export_data):
        tab_key  = data["tab_key"]
        schedule = data["schedule"]
        ranks    = data["ranks"]

        if idx > 0:
            doc.add_page_break()

        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = title.add_run("ΠΙΝΑΚΑΣ")
        set_font(r, 11, bold=True)
        r.underline = True

        # Subtitle
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        month_name_upper = GREEK_MONTHS_GEN[month].upper()
        r = subtitle.add_run(f"ΥΠΗΡΕΣΙΩΝ {TAB_TITLES.get(tab_key, tab_key)} ΜΗΝΟΣ {month_name_upper} {year}")
        set_font(r, 11, bold=True)

        doc.add_paragraph()  # blank line

        # Table
        table = doc.add_table(rows=days_in_month + 1, cols=4)
        table.style = "Table Grid"

        # Column widths via tblGrid
        tbl = table._tbl
        tblGrid = _OE("w:tblGrid")
        for w_cm in col_widths:
            gc = _OE("w:gridCol")
            gc.set(qn("w:w"), str(int(Cm(w_cm).twips)))
            tblGrid.append(gc)
        existing_grid = tbl.find(qn("w:tblGrid"))
        if existing_grid is not None:
            tbl.remove(existing_grid)
        tbl.insert(1, tblGrid)

        # Header row
        hdr_row = table.rows[0]
        set_row_height(hdr_row)
        hdr_labels = ["ΗΜΕΡΟΜΗΝΙΑ", "ΗΜΕΡΑ", "ΒΑΘΜΟΣ", "ΟΝΟΜΑΤΕΠΩΝΥΜΟ"]
        for ci, cell in enumerate(hdr_row.cells):
            set_cell_width(cell, col_widths[ci])
            set_cell_valign_center(cell)
            cell.text = ""
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = para.add_run(hdr_labels[ci])
            set_font(r, 11, bold=True)
            shading = parse_xml(r'<w:shd {} w:fill="D3D3D3"/>'.format(nsdecls("w")))
            cell._element.get_or_add_tcPr().append(shading)

        # Data rows
        for day in range(1, days_in_month + 1):
            row   = table.rows[day]
            set_row_height(row)
            cells = row.cells

            date_obj   = dt.date(year, month, day)
            weekday_gr = {
                "Monday":    "ΔΕΥΤΕΡΑ",
                "Tuesday":   "ΤΡΙΤΗ",
                "Wednesday": "ΤΕΤΑΡΤΗ",
                "Thursday":  "ΠΕΜΠΤΗ",
                "Friday":    "ΠΑΡΑΣΚΕΥΗ",
                "Saturday":  "ΣΑΒΒΑΤΟ",
                "Sunday":    "ΚΥΡΙΑΚΗ",
            }.get(date_obj.strftime("%A"), "")

            person   = schedule.get(day, "???")
            rank     = ranks.get(person, "")
            date_str = f"{day}-{GREEK_MONTH_ABBR[month].upper()}-{str(year)[-2:]}"
            row_vals = [date_str, weekday_gr, rank, person.upper()]

            for ci, cell in enumerate(cells):
                set_cell_width(cell, col_widths[ci])
                set_cell_valign_center(cell)
                cell.text = ""
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = para.add_run(row_vals[ci])
                set_font(r, 11)

            # Gray background for holidays
            if day_bucket(year, month, day, extra_holidays) == "HOLIDAY":
                for cell in cells:
                    shading = parse_xml(r'<w:shd {} w:fill="D3D3D3"/>'.format(nsdecls("w")))
                    cell._element.get_or_add_tcPr().append(shading)

    doc.save(filepath)