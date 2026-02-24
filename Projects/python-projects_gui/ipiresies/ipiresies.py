# ipiresies.py
# Scheduler με Two-Phase Backtracking Solver
# Optimized για ταχύτητα και Επιθυμίες

import calendar
import datetime as dt
import json
import sys, os
import random
import tkinter as tk
from tkinter import ttk, messagebox, filedialog

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls
from docx.oxml.ns import qn


# -----------------------------
# Constants
# -----------------------------
MIN_GAP_STRICT = 2
W_SCORE = 1.0
H_SCORE = 1.5

# Debug / tuning
DEBUG = False
SOLVER_MAX_TIME = 25.0  # seconds
SOLVER_MAX_RECURSION_MULT = 2  # recursion depth multiplier for weekdays phase

TAB_KEYS = ["AYDM", "BAYDM", "FKX", "PYLI"]
TAB_TITLES = {
    "AYDM": "ΑΥΔΜ",
    "BAYDM": "ΒΑΥΔΜ",
    "FKX": "ΦΚΧ",
    "PYLI": "ΠΥΛΗ",
}

GREEK_MONTHS_GEN = {
    1: "ΙΑΝΟΥΑΡΙΟΥ", 2: "ΦΕΒΡΟΥΑΡΙΟΥ", 3: "ΜΑΡΤΙΟΥ", 4: "ΑΠΡΙΛΙΟΥ",
    5: "ΜΑΪΟΥ", 6: "ΙΟΥΝΙΟΥ", 7: "ΙΟΥΛΙΟΥ", 8: "ΑΥΓΟΥΣΤΟΥ",
    9: "ΣΕΠΤΕΜΒΡΙΟΥ", 10: "ΟΚΤΩΒΡΙΟΥ", 11: "ΝΟΕΜΒΡΙΟΥ", 12: "ΔΕΚΕΜΒΡΙΟΥ",
}

GREEK_MONTH_ABBR = {
    1: "ΙΑΝ", 2: "ΦΕΒ", 3: "ΜΑΡ", 4: "ΑΠΡ", 5: "ΜΑΪ", 6: "ΙΟΥΝ",
    7: "ΙΟΥΛ", 8: "ΑΥΓ", 9: "ΣΕΠ", 10: "ΟΚΤ", 11: "ΝΟΕ", 12: "ΔΕΚ",
}


# -----------------------------
# Persistence
# -----------------------------
def get_base_dir() -> str:
    if getattr(sys, "frozen", False):
        return os.path.dirname(sys.executable)
    else:
        return os.path.dirname(os.path.abspath(__file__))

def get_state_path() -> str:
    """Αρχείο για τα inputs (άτομα, ρυθμίσεις) - autosave στο κλείσιμο."""
    return os.path.join(get_base_dir(), "scheduler_people.json")

def get_schedule_path() -> str:
    """Αρχείο για τα παραγόμενα προγράμματα - αποθηκεύεται μόνο χειροκίνητα."""
    return os.path.join(get_base_dir(), "scheduler_schedule.json")


def safe_int(s: str, default: int) -> int:
    try:
        return int(str(s).strip())
    except Exception:
        return default


# -----------------------------
# Core date helpers
# -----------------------------
def is_weekend(year: int, month: int, day: int) -> bool:
    wd = dt.date(year, month, day).weekday()
    return wd in (5, 6)


def parse_days_list(s: str, days_in_month: int) -> set[int]:
    s = (s or "").strip()
    if not s:
        return set()
    
    out: set[int] = set()
    parts = [p.strip() for p in s.split(",") if p.strip()]
    for p in parts:
        if "-" in p:
            a, b = p.split("-", 1)
            a, b = int(a), int(b)
            lo, hi = min(a, b), max(a, b)
            for d in range(lo, hi + 1):
                if 1 <= d <= days_in_month:
                    out.add(d)
                else:
                    raise ValueError(f"Ημέρα εκτός μήνα: {d}")
        else:
            d = int(p)
            if 1 <= d <= days_in_month:
                out.add(d)
            else:
                raise ValueError(f"Ημέρα εκτός μήνα: {d}")
    return out


def day_bucket(year: int, month: int, day: int, extra_holidays: set[int]) -> str:
    if day in extra_holidays:
        return "HOLIDAY"
    return "HOLIDAY" if is_weekend(year, month, day) else "WEEKDAY"



# -----------------------------
# Solver exceptions
# -----------------------------
class ScheduleError(RuntimeError):
    def __init__(self, message: str, details: str | None = None):
        super().__init__(message)
        self.details = details or message


# -----------------------------
# TOTAL quotas
# -----------------------------
def compute_total_quotas(names: list[str], days_in_month: int, rng: random.Random | None = None) -> dict[str, int]:
    if rng is None:
        import time; rng = random.Random(int(time.time() * 1000) % (2**32))
    n = len(names)
    base, extra = divmod(days_in_month, n)
    shuffled = names[:]
    rng.shuffle(shuffled)
    plus = set(shuffled[:extra])
    return {nm: (base + 1 if nm in plus else base) for nm in names}


def compute_total_quotas_with_history(
    names: list[str],
    ranks: dict[str, str],
    days_in_month: int,
    cumulative_stats: dict[str, dict],
    rng: random.Random,
) -> dict[str, int]:
    """
    Compute TOTAL service quotas with GROUP-BASED MIN-MAX balancing.
    
    Simple Logic:
    - People with LEAST cumulative total in their group → Get +1
    - People with MOST cumulative total in their group → Get base
    - Random tie-breaking for people with same totals
    """
    base, extra = divmod(days_in_month, len(names))
    quotas = {name: base for name in names}
    
    if extra == 0:
        return quotas
    
    # Group people by rank
    groups = {"OFFICERS": [], "OTHERS": []}
    for name in names:
        rank = ranks.get(name, "")
        group = get_person_group(rank)
        groups[group].append(name)
    
    remaining_extra = extra
    
    for group_name, group_members in groups.items():
        if not group_members or remaining_extra == 0:
            continue
        
        # Calculate totals for this group
        group_totals = {}
        for name in group_members:
            total = (
                cumulative_stats.get(name, {}).get("total_weekdays", 0) +
                cumulative_stats.get(name, {}).get("total_fridays", 0) +
                cumulative_stats.get(name, {}).get("total_holidays", 0)
            )
            group_totals[name] = total
        
        if not group_totals:
            continue
        
        # Sort by total WITH RANDOM TIE-BREAKING
        members_sorted = sorted(
            group_members,
            key=lambda n: (group_totals.get(n, 0), rng.random())  # Random tie-breaker!
        )
        
        # Give +1 to people with lowest totals
        group_extra = min(remaining_extra, len(group_members))
        for i in range(group_extra):
            quotas[members_sorted[i]] = base + 1
            remaining_extra -= 1
    
    # Distribute remaining
    if remaining_extra > 0:
        available = [n for n in names if quotas[n] == base]
        if available:
            # Sort with random tie-breaking
            available_sorted = sorted(
                available,
                key=lambda n: (
                    cumulative_stats.get(n, {}).get("total_weekdays", 0) +
                    cumulative_stats.get(n, {}).get("total_fridays", 0) +
                    cumulative_stats.get(n, {}).get("total_holidays", 0),
                    rng.random()
                )
            )
            for i in range(min(remaining_extra, len(available_sorted))):
                quotas[available_sorted[i]] = base + 1
    
    return quotas


def compute_holiday_quotas(names: list[str], total_holidays: int, rng: random.Random | None = None) -> dict[str, int]:
    """
    Distribute holiday services with ±1 rule.
    Same logic as total quotas but for holidays only.
    """
    if rng is None:
        import time; rng = random.Random(int(time.time() * 1000) % (2**32))
    n = len(names)
    base, extra = divmod(total_holidays, n)
    shuffled = names[:]
    rng.shuffle(shuffled)
    plus = set(shuffled[:extra])
    return {nm: (base + 1 if nm in plus else base) for nm in names}


def normalize_quotas_with_max(
    names: list[str],
    quotas_total: dict[str, int],
    days_in_month: int,
    max_caps: dict[str, int] | None,
) -> dict[str, int]:
    if not max_caps:
        return {nm: int(quotas_total.get(nm, 0)) for nm in names}
    
    q = {nm: int(quotas_total.get(nm, 0)) for nm in names}
    
    for nm in names:
        mx = max_caps.get(nm)
        if mx is None:
            continue
        try:
            mx_i = int(mx)
        except Exception:
            continue
        if mx_i > 0:
            q[nm] = min(q[nm], mx_i)
    
    remaining = days_in_month - sum(q.values())
    if remaining < 0:
        raise ScheduleError(
            "Τα MAX όρια οδηγούν σε υπερανάθεση.",
            details="Έλεγξε τα MAX."
        )
    
    def headroom(nm: str) -> int:
        mx = max_caps.get(nm)
        if mx is None:
            return 10**9
        try:
            mx_i = int(mx)
        except Exception:
            return 10**9
        if mx_i <= 0:
            return 10**9
        return mx_i - q[nm]
    
    while remaining > 0:
        candidates = [nm for nm in names if headroom(nm) > 0]
        if not candidates:
            raise ScheduleError(
                "Τα MAX όρια είναι πολύ χαμηλά.",
                details="Αυξησε τα MAX."
            )
        
        # IMPROVED: Prioritize people WITHOUT tight MAX constraints
        # If someone has MAX that's close to their current quota, they're constrained
        # Give the +1 to people with MORE headroom (less constrained)
        def priority(nm: str) -> tuple:
            current_headroom = headroom(nm)
            current_quota = q[nm]
            # Priority: (large headroom, low quota)
            # People with large headroom (e.g. no MAX or high MAX) should get +1 first
            return (-current_headroom, current_quota)
        
        candidates.sort(key=priority)
        chosen = candidates[0]
        q[chosen] += 1
        remaining -= 1
    
    return q


def compute_auto_max_from_leaves(
    names: list[str],
    days_in_month: int,
    leaves: dict[str, set[int]],
    min_gap: int = 2,
    log_cb=None,
) -> tuple[dict[str, int], dict[str, int]]:
    """
    AUTO-CALCULATE MAX based on available days.
    
    If someone has very limited available days (e.g., leave 5-31 = only 4 days available),
    calculate the maximum possible services they can do with MIN_GAP constraint.
    
    Also determines who needs MIN_GAP=1 for full utilization.
    
    Returns: 
        - auto_max: dict[name] -> max_services
        - custom_min_gap: dict[name] -> min_gap (1 for constrained people)
    """
    if log_cb is None:
        log_cb = lambda m: print(m)
    
    auto_max = {}
    custom_min_gap = {}
    
    for name in names:
        leave_days = leaves.get(name, set())
        available_days = [d for d in range(1, days_in_month + 1) if d not in leave_days]
        num_available = len(available_days)
        
        if num_available == 0:
            raise ScheduleError(
                f"Αδύνατο πρόγραμμα: {name}",
                details=f"Ο/Η {name} δεν έχει καμία διαθέσιμη μέρα (όλες άδειες)!"
            )
        
        # Calculate max possible services with MIN_GAP=2
        max_with_gap2 = (num_available + min_gap) // (min_gap + 1)
        
        # Calculate max possible services with MIN_GAP=1
        max_with_gap1 = (num_available + 1) // 2
        
        # NEW: Calculate OPTIMAL gap for best spread
        # Idea: If someone has 30 available days and needs 3 services,
        # use gap=14 to spread them evenly (days 1, 15, 29) instead of gap=2 (days 1, 3, 5)
        
        # THRESHOLD for when to apply custom gap
        # If someone has < 1/3 of month available, they're constrained
        GAP1_THRESHOLD = 3  # Use special handling if available < month/3
        
        if num_available < days_in_month / GAP1_THRESHOLD:
            # CONSTRAINED: Use minimum viable gap
            # Try GAP=2 first, then GAP=1 if needed
            if max_with_gap1 > max_with_gap2:
                auto_max[name] = max_with_gap1
                custom_min_gap[name] = 1  # Use GAP=1 for this person
                log_cb(f"  🔒 AUTO-MAX: {name} έχει μόνο {num_available} διαθέσιμες μέρες → MAX={max_with_gap1} (με GAP=1)")
            else:
                auto_max[name] = max_with_gap2
                custom_min_gap[name] = 2
                log_cb(f"  🔒 AUTO-MAX: {name} έχει μόνο {num_available} διαθέσιμες μέρες → MAX={max_with_gap2} (με GAP=2)")
        elif num_available < days_in_month / 2:
            # MEDIUM constraint: set MAX but use standard GAP=2
            auto_max[name] = max_with_gap2
            custom_min_gap[name] = 2
            log_cb(f"  🔒 AUTO-MAX: {name} έχει {num_available} διαθέσιμες μέρες → MAX={max_with_gap2} (με GAP=2)")
        else:
            # PLENTY of available days: do NOT increase GAP beyond the global setting.
            # A larger per-person GAP over-constrains the solver and can make feasible months unsatisfiable.
            custom_min_gap[name] = min_gap
    return auto_max, custom_min_gap



# -----------------------------
# PRIORITY ASSIGNMENT for MAX-constrained people
# -----------------------------
def assign_max_constrained_people(
    names: list[str],
    year: int,
    month: int,
    extra_holidays: set[int],
    leaves: dict[str, set[int]],
    max_caps: dict[str, int],
    custom_min_gap: dict[str, int] | None = None,
    log_cb=None,
) -> tuple[dict[int, str], dict[str, int]]:
    """
    Pre-assign people with explicit MAX constraints.
    
    NEW: Smart forced assignment for highly constrained people.
    If someone has very few available days and a tight MAX, 
    calculate the ONLY possible pattern and pre-assign it.
    
    Returns:
        - schedule: dict[day] -> person (pre-assigned days)
        - remaining_quotas: updated quotas after pre-assignment
    """
    if log_cb is None:
        log_cb = lambda m: print(m)
    
    if custom_min_gap is None:
        custom_min_gap = {}
    
    _, days_in_month = calendar.monthrange(year, month)
    all_days = list(range(1, days_in_month + 1))
    
    schedule = {}
    remaining_quotas = {}
    
    # Identify people WITH explicit MAX
    max_people = []
    normal_people = []
    
    for name in names:
        if name in max_caps and max_caps[name] is not None:
            try:
                max_val = int(max_caps[name])
                if max_val > 0:
                    max_people.append((name, max_val))
                    remaining_quotas[name] = 0  # Already handled
                    continue
            except:
                pass
        
        # Normal people - will be assigned later
        normal_people.append(name)
        remaining_quotas[name] = 0  # Will be calculated later
    
    if not max_people:
        log_cb("  ℹ️  Κανένα άτομο με MAX constraint")
        return {}, remaining_quotas
    
    log_cb(f"  🎯 ΦΑΣΗ 0: Ανάθεση MAX-constrained ατόμων")
    
    # Assign each MAX person
    for person, max_services in max_people:
        # Get this person's min_gap
        person_min_gap = custom_min_gap.get(person, 2)
        
        log_cb(f"    Ανάθεση {person}: MAX={max_services}, GAP={person_min_gap}")
        
        # Get available days for this person
        leave_days = leaves.get(person, set())
        available = [d for d in all_days if d not in leave_days and d not in schedule]
        
        if len(available) < max_services:
            raise ScheduleError(
                f"Αδύνατο MAX για {person}",
                details=f"{person} έχει MAX={max_services} αλλά μόνο {len(available)} διαθέσιμες μέρες!"
            )
        
        # SMART ASSIGNMENT: Calculate forced pattern with proper gap
        # Use greedy algorithm with min_gap constraint
        forced_pattern = []
        for day in available:
            # Can we add this day?
            ok = True
            for assigned_day in forced_pattern:
                if abs(day - assigned_day) <= person_min_gap:
                    ok = False
                    break
            
            if ok:
                forced_pattern.append(day)
                if len(forced_pattern) == max_services:
                    break
        
        if len(forced_pattern) == max_services:
            log_cb(f"      ✅ Pattern με GAP={person_min_gap}: {forced_pattern}")
            assigned_days = forced_pattern
        else:
            # Couldn't find enough days with the gap - this should not happen if AUTO-MAX is correct
            raise ScheduleError(
                f"Αδύνατο pattern για {person}",
                details=f"{person} έχει MAX={max_services}, GAP={person_min_gap} αλλά δεν χωράνε στις {len(available)} διαθέσιμες μέρες!"
            )
        
        for day in assigned_days:
            schedule[day] = person
    
    # Calculate remaining quotas for normal people
    days_used = len(schedule)
    days_remaining = days_in_month - days_used
    
    if normal_people:
        base, extra = divmod(days_remaining, len(normal_people))
        import random
        import time; rng = random.Random(int(time.time() * 1000) % (2**32))
        shuffled = normal_people[:]
        rng.shuffle(shuffled)
        plus = set(shuffled[:extra])
        
        for name in normal_people:
            remaining_quotas[name] = base + (1 if name in plus else 0)
    
    log_cb(f"    ✅ Pre-assigned: {days_used} μέρες")
    log_cb(f"    📊 Υπόλοιπα quotas: {remaining_quotas}")
    
    return schedule, remaining_quotas


# -----------------------------
# TWO-PHASE Backtracking (Holidays first, then weekdays)
# -----------------------------
def solve_with_two_phase_backtracking(
    names: list[str],
    year: int,
    month: int,
    extra_holidays: set[int],
    leaves: dict[str, set[int]],
    quotas: dict[str, int],
    holiday_quotas: dict[str, int],
    friday_quotas: dict[str, int] | None = None,
    preferences: dict[str, set[int]] | None = None,
    min_gap: int = 2,
    custom_min_gap: dict[str, int] | None = None,
    pre_assigned: dict[int, str] | None = None,
    log_cb=None,
    ignore_weekend_pair_days: set[int] | None = None,  # days where Sat/Fri rule is relaxed
) -> tuple[dict[int, str], dict, str]:
    """
    Two-phase solver with score balancing and per-person min_gap support.
    
    PHASE 1: Assign all holidays first (easier, fewer days)
    PHASE 2: Assign remaining weekdays (with score balancing)
    
    custom_min_gap: dict[name] -> min_gap override for specific people
    pre_assigned: dict[day] -> person (already assigned, skip these days)
    preferences: dict[name] -> set of preferred days (soft constraint)
    """
    if log_cb is None:
        log_cb = lambda m: print(m)
    
    if preferences is None:
        preferences = {}
    
    if custom_min_gap is None:
        custom_min_gap = {}
    
    if pre_assigned is None:
        pre_assigned = {}
    
    if friday_quotas is None:
        friday_quotas = {p: 0 for p in names}

    # Flag to enable/disable friday quota enforcement (may be relaxed if solver fails)
    enforce_friday_quota = [True]

    log_cb("🎯 Two-phase solver (Αργίες πρώτα, μετά καθημερινές)...")
    
    _, days_in_month = calendar.monthrange(year, month)
    all_days = list(range(1, days_in_month + 1))
    
    # Filter out pre-assigned days
    available_days = [d for d in all_days if d not in pre_assigned]
    
    holiday_days = sorted([d for d in available_days if day_bucket(year, month, d, extra_holidays) == "HOLIDAY"])
    weekdays = sorted([d for d in available_days if day_bucket(year, month, d, extra_holidays) == "WEEKDAY"])
    friday_days = set(d for d in weekdays if dt.date(year, month, d).weekday() == 4)
    
    log_cb(f"  📅 Pre-assigned: {len(pre_assigned)} μέρες")
    log_cb(f"  📅 Αργίες: {len(holiday_days)} μέρες")
    log_cb(f"  📅 Καθημερινές: {len(weekdays)} μέρες")
    
    # Track scores for balancing
    current_scores = {p: 0.0 for p in names}
    
    # Helper to get effective min_gap for a person
    def get_min_gap(person: str) -> int:
        return custom_min_gap.get(person, min_gap)
    
    # Helper to check weekend-specific forbidden pairs
    def is_forbidden_weekend_pair(day1: int, day2: int) -> bool:
        """
        Service-specific Fri/Sat gap rules (FORBIDDEN if next assignment is too close):

        AYDM & PYLI:
          - Friday requires gap >= 3
          - Saturday requires gap >= 4

        FKX & BAYDM:
          - Friday requires gap >= 4
          - Saturday requires gap >= 5
        """
        if day1 >= day2:
            return False

        gap_days = day2 - day1

        # Detect service/tab key from enclosing scope (tab_key exists in your program)
        try:
            service = tab_key  # e.g. "AYDM", "BAYDM", "FKX", "PYLI"
        except NameError:
            service = "AYDM"   # safe default if not in scope

        # Map per-service rules
        if service in ("AYDM", "PYLI"):
            req_gap_fri = 3
            req_gap_sat = 4
        elif service in ("FKX", "BAYDM"):
            req_gap_fri = 4
            req_gap_sat = 5
        else:
            # Unknown service -> no special weekend rule
            return False

        # Determine weekday of day1
        try:
            weekday1 = dt.date(year, month, day1).weekday()  # Mon=0 ... Fri=4 Sat=5 Sun=6
        except ValueError:
            return False

        # Apply rule only if day1 is Friday or Saturday
        if weekday1 == 4:  # Friday
            return gap_days <= req_gap_fri
        if weekday1 == 5:  # Saturday
            return gap_days <= req_gap_sat

        return False

    
    # Helper to check weekend-specific constraints
    def violates_weekend_rule(person: str, day1: int, day2: int) -> bool:
        """
        Check if assigning person to both day1 and day2 violates weekend rules:
        - Friday (day1) -> Wednesday next week (day2): NOT ALLOWED (gap=5)
        - Saturday (day1) -> Thursday next week (day2): NOT ALLOWED (gap=5)
        """
        if day1 >= day2:  # Only check forward
            return False
        
        gap = day2 - day1
        
        # Get weekday for day1
        date1 = dt.date(year, month, day1)
        weekday1 = date1.weekday()  # Monday=0, Sunday=6
        
        # Friday (4) + 6 days = Thursday next week
        if weekday1 == 4 and gap == 5:  # Friday -> 5 days later (Wednesday)
            return True
        
        # Saturday (5) + 5 days = Thursday next week  
        if weekday1 == 5 and gap == 5:  # Saturday -> 5 days later (Thursday)
            return True
        
        return False
    
    # ==================== PHASE 1: HOLIDAYS ====================
    log_cb("  🎉 ΦΑΣΗ 1: Ανάθεση αργιών...")
    
    def candidates_for_holiday(d: int, partial: dict[int, str]) -> list[str]:
        """Find who can work on holiday day d."""
        cands = []
        for p in names:
            # Check leave
            if d in leaves.get(p, set()):
                continue
            
            # Check holiday quota
            holiday_used = sum(1 for dd, pp in partial.items() if pp == p and dd in holiday_days)
            if holiday_used >= holiday_quotas[p]:
                continue
            
            # Check min gap (with per-person gap) - includes pre-assigned
            person_gap = get_min_gap(p)
            full_assigned = {**pre_assigned, **partial}
            ok = True
            for offset in range(-person_gap, person_gap + 1):
                if offset == 0:
                    continue
                check_day = d + offset
                if check_day in full_assigned and full_assigned[check_day] == p:
                    # If this conflict comes only from a FORCED (pre-assigned) preference,
                    # do NOT block the candidate. Forced preference days are "locked" and
                    # excluded from GAP constraints for additional assignments.
                    if check_day in pre_assigned:
                        continue
                    ok = False
                    break
            
            if not ok:
                continue
            
            # Check weekend-specific forbidden pairs (includes pre-assigned)
            if ignore_weekend_pair_days is None or d not in ignore_weekend_pair_days:
                for other_day, other_person in full_assigned.items():
                    if other_day in pre_assigned:
                        continue
                    if other_person == p:
                        if is_forbidden_weekend_pair(other_day, d) or is_forbidden_weekend_pair(d, other_day):
                            ok = False
                            break

            if ok:
                cands.append(p)
        
        return cands
    
    # Sort holidays by difficulty (with EXTREME priority for forced moves)
    def holiday_difficulty(d: int) -> int:
        num_candidates = len(candidates_for_holiday(d, {}))
        # FORCED MOVE (only 1 candidate) → HIGHEST PRIORITY!
        if num_candidates == 1:
            return -1000
        # NO CANDIDATES → Should fail immediately, but mark as very hard
        if num_candidates == 0:
            return -999
        return num_candidates
    
    sorted_holidays = sorted(holiday_days, key=holiday_difficulty)
    
    # Recursion depth limit (VERY AGGRESSIVE - 1x only!)
    max_recursion_depth_holidays = len(sorted_holidays)
    recursion_cutoff_holidays = [0]
    
    def backtrack_holidays(idx: int, schedule: dict[int, str], remaining_hol: dict[str, int], depth: int = 0) -> bool:
        # VERY AGGRESSIVE Recursion depth limit
        if depth > max_recursion_depth_holidays:
            recursion_cutoff_holidays[0] += 1
            if recursion_cutoff_holidays[0] == 1:
                log_cb(f"  ⚠️  Holiday backtracking stuck, aborting...")
            return False
        
        if idx == len(sorted_holidays):
            return True
        
        # FAIL-FAST #1: Total check
        remaining_holidays_count = len(sorted_holidays) - idx
        total_remaining_hol_quota = sum(remaining_hol.values())
        if total_remaining_hol_quota > remaining_holidays_count:
            return False
        
        # FAIL-FAST #2: Individual feasibility
        for p in names:
            if remaining_hol[p] <= 0:
                continue
            
            possible_holidays = 0
            for future_idx in range(idx, len(sorted_holidays)):
                future_day = sorted_holidays[future_idx]
                if future_day not in leaves.get(p, set()):
                    possible_holidays += 1
            
            if possible_holidays < remaining_hol[p]:
                return False
        
        d = sorted_holidays[idx]
        cands = candidates_for_holiday(d, schedule)
        
        if not cands:
            return False
        
        # PRIORITY SORTING with PREFERENCES FIRST + RANDOMIZATION for ties!
        def holiday_priority(p: str) -> tuple:
            has_preference = 1 if d in preferences.get(p, set()) else 0
            return (
                has_preference,         # ΕΠΙΘΥΜΙΑ - HIGHEST!
                remaining_hol[p],       # Remaining holiday quota
                -current_scores[p],     # Score balancing
                random.random(),        # RANDOMIZATION for ties!
            )
        
        cands.sort(key=holiday_priority, reverse=True)
        
        for p in cands:
            schedule[d] = p
            remaining_hol[p] -= 1
            current_scores[p] += H_SCORE
            
            if backtrack_holidays(idx + 1, schedule, remaining_hol, depth + 1):
                return True
            
            del schedule[d]
            remaining_hol[p] += 1
            current_scores[p] -= H_SCORE
        
        return False
    
    # Try to assign holidays
    holiday_schedule = {}
    remaining_holiday_quotas = holiday_quotas.copy()
    
    if not backtrack_holidays(0, holiday_schedule, remaining_holiday_quotas, depth=0):
        # Provide detailed error info
        error_details = f"Δεν μπορούν να κατανεμηθούν οι αργίες.\n"
        error_details += f"Αργίες: {len(holiday_days)} μέρες\n"
        error_details += f"Holiday quotas: {holiday_quotas}\n"
        
        # Check which days are problematic
        for d in sorted_holidays[:5]:  # First 5 difficult days
            cands = candidates_for_holiday(d, {})
            date = dt.date(year, month, d)
            weekday = date.strftime("%A")
            error_details += f"  Μέρα {d} ({weekday}): {len(cands)} υποψήφιοι - {cands}\n"
        
        raise ScheduleError(
            "Αδύνατη κατανομή αργιών",
            details=error_details
        )
    
    log_cb(f"  ✅ Αργίες ανατέθηκαν επιτυχώς!")
    
    # ==================== PHASE 2: WEEKDAYS ====================
    log_cb("  📋 ΦΑΣΗ 2: Ανάθεση καθημερινών (με score balancing)...")
    
    # Calculate remaining total quotas after holidays
    remaining_total_quotas = quotas.copy()
    for d, p in holiday_schedule.items():
        remaining_total_quotas[p] -= 1
    
    log_cb(f"  📊 Υπόλοιπες υπηρεσίες: {remaining_total_quotas}")
    log_cb(f"  📊 Scores μετά τις αργίες: {current_scores}")
    
    def candidates_for_weekday(d: int, partial: dict[int, str]) -> list[str]:
        """Find who can work on weekday d."""
        cands = []
        for p in names:
            # Check leave
            if d in leaves.get(p, set()):
                continue
            
            # Check remaining total quota
            weekday_used = sum(1 for dd, pp in partial.items() if pp == p and dd in weekdays)
            if weekday_used >= remaining_total_quotas[p]:
                continue
            
            # Check friday quota (only when day d is a Friday, and only if enforced)
            if enforce_friday_quota[0] and d in friday_days:
                friday_used = sum(1 for dd, pp in partial.items() if pp == p and dd in friday_days)
                if friday_used >= friday_quotas.get(p, 0):
                    continue
            
            # Check min gap (with ALL assigned days: pre-assigned + holidays + weekdays, per-person gap)
            full_schedule = {**pre_assigned, **holiday_schedule, **partial}
            person_gap = get_min_gap(p)
            ok = True
            for offset in range(-person_gap, person_gap + 1):
                if offset == 0:
                    continue
                check_day = d + offset
                if check_day in full_schedule and full_schedule[check_day] == p:
                    # Ignore GAP conflicts against FORCED (pre-assigned) preference days.
                    if check_day in pre_assigned:
                        continue
                    ok = False
                    break
            
            if not ok:
                continue
            
            # Check weekend-specific forbidden pairs
            if ignore_weekend_pair_days is None or d not in ignore_weekend_pair_days:
                for other_day, other_person in full_schedule.items():
                    if other_day in pre_assigned:
                        continue
                    if other_person == p:
                        if is_forbidden_weekend_pair(other_day, d) or is_forbidden_weekend_pair(d, other_day):
                            ok = False
                            break

            if ok:
                cands.append(p)
        
        return cands
    
    # Sort weekdays by difficulty (FORCED MOVES = HIGHEST PRIORITY!)
    # IMPORTANT: Consider pre-assigned days when calculating difficulty!
    def weekday_difficulty(d: int) -> int:
        num_candidates = len(candidates_for_weekday(d, {}))  # Empty schedule for initial sort
        # FORCED MOVE (only 1 candidate) → HIGHEST PRIORITY!
        if num_candidates == 1:
            return -1000
        if num_candidates == 0:
            return -999
        return num_candidates
    
    sorted_weekdays = sorted(weekdays, key=weekday_difficulty)
    
    # Timeout tracking (10 seconds max!)
    import time
    start_time = time.time()
    max_time = SOLVER_MAX_TIME  # seconds
    
    # Recursion depth tracking (VERY AGGRESSIVE LIMIT!)
    # Only allow 1x the number of days - if it needs more, it's stuck
    max_recursion_depth = len(sorted_weekdays) * SOLVER_MAX_RECURSION_MULT
    recursion_cutoff_count = [0]  # Track how many times we hit the limit
    
    def backtrack_weekdays(idx: int, schedule: dict[int, str], remaining: dict[str, int], depth: int = 0) -> bool:
        # TIMEOUT CHECK
        if time.time() - start_time > max_time:
            log_cb(f"  ⚠️  Timeout after {max_time}s, aborting backtracking...")
            return False
        
        # VERY AGGRESSIVE Recursion depth limit
        if depth > max_recursion_depth:
            recursion_cutoff_count[0] += 1
            if recursion_cutoff_count[0] == 1:  # Log only first time
                log_cb(f"  ⚠️  Backtracking stuck (depth > {max_recursion_depth}), aborting...")
            return False
        
        if idx == len(sorted_weekdays):
            return True
        
        # FAIL-FAST #1: Total quota check
        remaining_days_count = len(sorted_weekdays) - idx
        total_remaining_quota = sum(remaining.values())
        if total_remaining_quota > remaining_days_count:
            return False
        
        # FAIL-FAST #2: Individual feasibility check
        # For each person with remaining quota, check if they can possibly fulfill it
        for p in names:
            if remaining[p] <= 0:
                continue
            
            # Count how many of the remaining days this person can potentially work
            possible_days = 0
            for future_idx in range(idx, len(sorted_weekdays)):
                future_day = sorted_weekdays[future_idx]
                # Quick check: can this person work this day? (ignoring other assignments)
                if future_day not in leaves.get(p, set()):
                    possible_days += 1
            
            # If they can't possibly fulfill their quota, fail fast!
            if possible_days < remaining[p]:
                return False
        
        d = sorted_weekdays[idx]
        cands = candidates_for_weekday(d, schedule)
        
        if not cands:
            return False
        
        # PRIORITY SORTING with RANDOMIZATION for ties:
        # 1. PREFERENCE MATCH (highest priority!)
        # 2. Remaining quota (need to assign)
        # 3. Score balancing (fairness)
        # 4. RANDOM (for ties!)
        def candidate_priority(p: str) -> tuple:
            # Check if person has preference for this specific day
            has_preference = 1 if d in preferences.get(p, set()) else 0
            return (
                has_preference,        # ΕΠΙΘΥΜΙΑ - HIGHEST PRIORITY!
                remaining[p],          # Still need to assign (higher = better)
                -current_scores[p],    # Current score (lower = better)
                random.random(),       # RANDOMIZATION for ties!
            )
        
        cands.sort(key=candidate_priority, reverse=True)
        
        for p in cands:
            schedule[d] = p
            remaining[p] -= 1
            current_scores[p] += W_SCORE
            
            if backtrack_weekdays(idx + 1, schedule, remaining, depth + 1):
                return True
            
            del schedule[d]
            remaining[p] += 1
            current_scores[p] -= W_SCORE
        
        return False
    
    # Try to assign weekdays
    # Attempt 1: with friday quota strictly enforced
    weekday_schedule = {}
    remaining_weekday_quotas = remaining_total_quotas.copy()
    weekday_success = backtrack_weekdays(0, weekday_schedule, remaining_weekday_quotas, depth=0)

    # Attempt 2: if failed, relax friday quota constraint and retry
    if not weekday_success:
        enforce_friday_quota[0] = False
        recursion_cutoff_count[0] = 0
        weekday_schedule = {}
        remaining_weekday_quotas = remaining_total_quotas.copy()
        log_cb("  ⚠️  Friday quota constraint χαλάρωσε - επαναπροσπάθεια...")
        weekday_success = backtrack_weekdays(0, weekday_schedule, remaining_weekday_quotas, depth=0)

    if not weekday_success:
        # Provide detailed error info
        error_details = f"Οι αργίες ανατέθηκαν, αλλά δεν μπορούν να κατανεμηθούν οι καθημερινές.\n"
        error_details += f"Καθημερινές: {len(weekdays)} μέρες\n"
        error_details += f"Υπόλοιπα quotas: {remaining_total_quotas}\n"
        
        if recursion_cutoff_count[0] > 0:
            error_details += f"\n⚠️  Recursion limit hit {recursion_cutoff_count[0]} times (max depth: {max_recursion_depth})\n"
            error_details += f"Αυτό σημαίνει ότι το πρόβλημα είναι ΠΟΛΥ δύσκολο με τους τρέχοντες constraints.\n"
        
        # Check which days are problematic
        error_details += f"\n🔍 Πιο δύσκολες μέρες:\n"
        for d in sorted_weekdays[:5]:  # First 5 difficult days
            cands = candidates_for_weekday(d, {})
            date = dt.date(year, month, d)
            weekday = date.strftime("%A")
            error_details += f"  Μέρα {d} ({weekday}): {len(cands)} υποψήφιοι"
            if len(cands) <= 3:
                error_details += f" - {cands}"
            error_details += "\n"
        
        # Check for people with very limited availability
        error_details += f"\n👥 Άτομα με περιορισμένες μέρες:\n"
        for name in names:
            leave_days = leaves.get(name, set())
            available = len([d for d in weekdays if d not in leave_days])
            if available < len(weekdays) / 2:  # Less than half available
                error_details += f"  {name}: {available}/{len(weekdays)} διαθέσιμες καθημερινές\n"
        
        raise ScheduleError(
            "Αδύνατη κατανομή καθημερινών",
            details=error_details
        )
    
    log_cb(f"  ✅ Καθημερινές ανατέθηκαν επιτυχώς!")
    
    # Check final score spread
    final_scores = list(current_scores.values())
    spread = max(final_scores) - min(final_scores)
    log_cb(f"  📊 Final score spread: {spread:.2f}")
    
    # Combine ALL schedules: pre-assigned + holidays + weekdays
    final_schedule = {**pre_assigned, **holiday_schedule, **weekday_schedule}
    
    # Check preference matches
    if preferences:
        total_prefs = sum(len(days) for days in preferences.values())
        matched_prefs = 0
        for day, person in final_schedule.items():
            if day in preferences.get(person, set()):
                matched_prefs += 1
        
        if total_prefs > 0:
            match_pct = (matched_prefs / total_prefs) * 100
            log_cb(f"  💡 Επιθυμίες: {matched_prefs}/{total_prefs} ικανοποιήθηκαν ({match_pct:.1f}%)")
    
    # Build min_gap info for statistics
    gap_info = {}
    for name in names:
        gap_info[name] = get_min_gap(name)
    
    solve_info = {
        'solver': 'Two-Phase Backtracking',
        'pre_assigned': len(pre_assigned),
        'holidays_assigned': len(holiday_schedule),
        'weekdays_assigned': len(weekday_schedule),
        'score_spread': spread,
        'min_gap_used': gap_info,  # NEW: gap per person
    }
    
    return final_schedule, solve_info, 'Two-Phase'


# -----------------------------
# Original Backtracking Fallback (for relaxed constraints)


# -----------------------------
# UNIFIED Multi-Tab Scheduling
# -----------------------------
def solve_all_tabs_unified(
    all_tabs_data: dict[str, dict],
    year: int,
    month: int,
    extra_holidays: set[int],
    log_cb=None,
) -> dict[str, tuple]:
    """
    Solve all tabs SEQUENTIALLY to avoid same person on same day across tabs.
    
    Uses ONLY history for balancing (no running totals within month).
    Each tab independently balances based on cumulative history.
    
    Returns: {
        "AYDM": (schedule, quotas, meta, solve_info),
        ...
    }
    """
    if log_cb is None:
        log_cb = lambda *args, **kwargs: None

    # Support two styles:
    #   log_cb(message)
    #   log_cb(tab_key, message)   -> routes logs per tab in the GUI
    def _emit(tab_key, message: str):
        try:
            log_cb(tab_key, message)
        except TypeError:
            log_cb(message)

    _emit(None, "="*60)
    _emit(None, "🌐 UNIFIED SCHEDULING - Διαδοχική επεξεργασία")
    _emit(None, "="*60)
    
    results = {}
    global_assignments = {}  # {person: set of days already assigned}
    
    # Process each tab in order (ΑΥΔΜ, ΒΑΥΔΜ, ΦΚΧ, ΠΥΛΗ)
    for tab_key in TAB_KEYS:
        if tab_key not in all_tabs_data:
            continue
        
        tab_data = all_tabs_data[tab_key]
        names = tab_data["names"]
        leaves = tab_data["leaves"]
        ranks = tab_data["ranks"]
        max_caps = tab_data["max_caps"]
        preferences = tab_data.get("preferences", {})  # NEW: get preferences
        
        if not names:
            continue
        
        _emit(tab_key, f"\n📋 {TAB_TITLES[tab_key]}...")
        
        # Add previous assignments to leaves (cross-tab conflict prevention)
        extended_leaves = {}
        for person in names:
            person_leaves = set(leaves.get(person, set()))
            
            if person in global_assignments:
                already_assigned = global_assignments[person]
                person_leaves.update(already_assigned)
                _emit(tab_key, f"  ⚠️  {person} ήδη ανατεθειμένος σε: {sorted(already_assigned)}")
            
            extended_leaves[person] = person_leaves
        
        # Solve this tab (uses history only, no running totals)
        try:
            schedule, quotas, meta, solve_info = solve_schedule_best_effort(
                names=names,
                year=year,
                month=month,
                extra_holidays=extra_holidays,
                leaves=extended_leaves,
                preferences=preferences,  # NEW: pass preferences
                max_caps=max_caps,
                ranks=ranks,
                tab_key=tab_key,
                log_cb=lambda m, tk=tab_key: _emit(tk, f"    {m}"),
            )
            
            # Update global assignments
            for day, person in schedule.items():
                if person not in global_assignments:
                    global_assignments[person] = set()
                global_assignments[person].add(day)
            
            results[tab_key] = (schedule, quotas, meta, solve_info)
            _emit(tab_key, f"  ✅ {TAB_TITLES[tab_key]} ολοκληρώθηκε")
            
        except ScheduleError as e:
            _emit(tab_key, f"  ❌ {TAB_TITLES[tab_key]} αποτυχία: {e.message}")
            raise ScheduleError(
                f"Αποτυχία στο {TAB_TITLES[tab_key]}",
                details=e.details
            )
    
    _emit(None, "\n" + "="*60)
    _emit(None, "✅ Όλα τα προγράμματα ολοκληρώθηκαν!")
    _emit(None, "="*60)
    
    return results


# -----------------------------

# -----------------------------
# History Management
# -----------------------------
# Save history in the same directory as the exe/script
HISTORY_FILE = os.path.join(get_base_dir(), "ipiresies_history.json")


def get_person_group(rank: str) -> str:
    """
    Determine which group a person belongs to based on their rank.
    
    Group A (OFFICERS): Ranks ending in "ΛΓΟΣ" (ΛΓΟΣ, ΥΠΛΓΟΣ, ΑΝΘΛΓΟΣ)
    Group B (OTHERS): All other ranks (ΑΛΧΙΑΣ, ΕΠΧΙΑΣ, etc.)
    
    Returns: "OFFICERS" or "OTHERS"
    """
    rank_upper = rank.upper().strip()
    if rank_upper.endswith("ΛΓΟΣ"):
        return "OFFICERS"
    else:
        return "OTHERS"


def load_history() -> dict:
    """Load history from JSON file."""
    if not os.path.exists(HISTORY_FILE):
        return {}
    
    try:
        with open(HISTORY_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        print(f"Warning: Could not load history: {e}")
        return {}


def save_history(history: dict):
    """Save history to JSON file."""
    try:
        with open(HISTORY_FILE, 'w', encoding='utf-8') as f:
            json.dump(history, f, indent=2, ensure_ascii=False)
    except Exception as e:
        print(f"Error saving history: {e}")


def add_to_history(
    year: int,
    month: int,
    tab_key: str,
    schedule: dict[int, str],
    ranks: dict[str, str],
    extra_holidays: set[int],
):
    """
    Add a finalized schedule to history.
    
    NEW STRUCTURE (aggregated per person WITH RANK):
    {
      "2026": {
        "1": {
          "Γιάννης": {
            "rank": "ΛΓΟΣ",      // NEW: store rank!
            "weekday": 5,
            "friday": 1,
            "holiday": 2,
          }
        }
      }
    }
    """
    history = load_history()
    
    year_str = str(year)
    month_str = str(month)
    
    if year_str not in history:
        history[year_str] = {}
    if month_str not in history[year_str]:
        history[year_str][month_str] = {}
    
    # Αν το tab έχει ήδη αποθηκευτεί για αυτόν τον μήνα, σβήνουμε τα παλιά δεδομένα
    # και τα αντικαθιστούμε με τα νέα (αντί να τα αθροίζουμε)
    saved_tabs = history[year_str][month_str].get("_saved_tabs", [])
    if tab_key in saved_tabs:
        # Βρες και διέγραψε τις εγγραφές που ανήκουν σε αυτό το tab
        # (δεν ξέρουμε ποιοι ανήκουν σε ποιο tab, οπότε σβήνουμε όσα είναι στο schedule)
        people_in_schedule = set(schedule.values())
        for person in people_in_schedule:
            if person in history[year_str][month_str]:
                del history[year_str][month_str][person]

    # Count by category for each person
    for day, person in schedule.items():
        if person not in history[year_str][month_str]:
            history[year_str][month_str][person] = {
                "rank": ranks.get(person, ""),  # NEW: store rank
                "weekday": 0,
                "friday": 0,
                "holiday": 0,
            }
        
        # Determine category
        date_obj = dt.date(year, month, day)
        weekday_num = date_obj.weekday()  # 0=Mon, 4=Fri, 5=Sat, 6=Sun
        bucket = day_bucket(year, month, day, extra_holidays)
        
        if bucket == "HOLIDAY":
            history[year_str][month_str][person]["holiday"] += 1
        elif weekday_num == 4:  # Friday
            history[year_str][month_str][person]["friday"] += 1
        else:  # Mon-Thu
            history[year_str][month_str][person]["weekday"] += 1
    
    # Σημείωσε το tab ως αποθηκευμένο για αυτόν τον μήνα
    if "_saved_tabs" not in history[year_str][month_str]:
        history[year_str][month_str]["_saved_tabs"] = []
    if tab_key not in history[year_str][month_str]["_saved_tabs"]:
        history[year_str][month_str]["_saved_tabs"].append(tab_key)
    
    save_history(history)


def calculate_cumulative_stats(
    year: int,
    month: int,
    tab_key: str,  # Not used anymore but kept for compatibility
    names: list[str],
) -> dict[str, dict]:
    """
    Calculate cumulative stats for each person from start of year until (excluding) current month.
    
    NEW STRUCTURE: Aggregated (no per-service breakdown)
    
    Returns: {
        "Γιάννης": {
            "months": 2,
            "total_weekdays": 10,  // Mon-Thu only
            "total_fridays": 2,    // Fridays only
            "total_holidays": 4,   // Sat/Sun + extra
            "cumulative_score": 5.5,
        }
    }
    """
    history = load_history()
    year_str = str(year)
    
    cumulative = {}
    for name in names:
        cumulative[name] = {
            "months": 0,
            "total_weekdays": 0,
            "total_fridays": 0,
            "total_holidays": 0,
            "cumulative_score": 0.0,
        }
    
    # Check if we have history for this year
    if year_str not in history:
        return cumulative
    
    # Sum up all months before current month
    for past_month in range(1, month):
        month_str = str(past_month)
        if month_str not in history[year_str]:
            continue
        
        month_data = history[year_str][month_str]
        
        for name in names:
            if name in month_data:
                data = month_data[name]
                cumulative[name]["months"] += 1
                cumulative[name]["total_weekdays"] += data.get("weekday", 0)
                cumulative[name]["total_fridays"] += data.get("friday", 0)
                cumulative[name]["total_holidays"] += data.get("holiday", 0)
                
                # Calculate score: weekday=1, friday=1, holiday=1.5
                score = (data.get("weekday", 0) * W_SCORE + 
                        data.get("friday", 0) * W_SCORE +
                        data.get("holiday", 0) * H_SCORE)
                cumulative[name]["cumulative_score"] += score
    
    return cumulative


def add_current_month_stats(
    cumulative: dict[str, dict],
    schedules: dict[str, dict[int, str]],
    year: int,
    month: int,
    extra_holidays: set[int],
) -> dict[str, dict]:
    """
    Add current month's assignments (from multiple tabs) to cumulative stats.
    
    This is used during unified scheduling to track running totals across tabs.
    
    Args:
        cumulative: Existing cumulative stats from history
        schedules: {tab_key: {day: person}} - current month assignments
        year, month: Current scheduling period
        extra_holidays: Extra holiday dates
    
    Returns:
        Updated cumulative stats including current month
    """
    import datetime as dt
    
    result = {}
    for name, stats in cumulative.items():
        result[name] = stats.copy()
    
    # Add current month assignments
    for tab_key, schedule in schedules.items():
        for day, person in schedule.items():
            if person not in result:
                result[person] = {
                    "months": 0,
                    "total_weekdays": 0,
                    "total_fridays": 0,
                    "total_holidays": 0,
                    "cumulative_score": 0.0,
                }
            
            # Determine category
            date_obj = dt.date(year, month, day)
            weekday_num = date_obj.weekday()
            bucket = day_bucket(year, month, day, extra_holidays)
            
            if bucket == "HOLIDAY":
                result[person]["total_holidays"] += 1
            elif weekday_num == 4:  # Friday
                result[person]["total_fridays"] += 1
            else:  # Mon-Thu
                result[person]["total_weekdays"] += 1
    
    return result


def compute_holiday_quotas_with_history(
    names: list[str],
    total_holidays: int,
    cumulative_stats: dict[str, dict],
    ranks: dict[str, str],
    rng: random.Random,
    leaves: dict[str, set[int]] = None,  # NEW: to check available holidays
    year: int = None,  # NEW
    month: int = None,  # NEW
    extra_holidays: set[int] = None,  # NEW
) -> dict[str, int]:
    """
    Compute holiday quotas with GROUP-BASED MIN-MAX balancing.
    
    Simple Logic:
    - People with FEWEST holidays in their group → Get +1
    - People with MOST holidays in their group → Get base
    - Random tie-breaking
    - ZERO quota for people with NO available holidays
    """
    if leaves is None:
        leaves = {}
    if extra_holidays is None:
        extra_holidays = set()
    
    # First, check which people have available holidays
    available_for_holidays = []
    if year and month:
        import calendar as cal
        _, days_in_month = cal.monthrange(year, month)
        for name in names:
            leave_days = leaves.get(name, set())
            # Check if person has at least one available holiday
            has_holiday = False
            for d in range(1, days_in_month + 1):
                if d not in leave_days and day_bucket(year, month, d, extra_holidays) == "HOLIDAY":
                    has_holiday = True
                    break
            if has_holiday:
                available_for_holidays.append(name)
    else:
        # Fallback: assume everyone available
        available_for_holidays = names[:]
    
    # Initialize quotas: ZERO for people without holidays, base for others
    if not available_for_holidays:
        # Nobody has holidays available - zero for all
        return {name: 0 for name in names}
    
    base, extra = divmod(total_holidays, len(available_for_holidays))
    quotas = {name: 0 for name in names}  # Start with zero
    for name in available_for_holidays:
        quotas[name] = base
    
    if extra == 0:
        return quotas
    
    # Group people by rank
    groups = {"OFFICERS": [], "OTHERS": []}
    # Group ONLY people who have available holidays
    groups = {"OFFICERS": [], "OTHERS": []}
    for name in available_for_holidays:
        rank = ranks.get(name, "")
        group = get_person_group(rank)
        groups[group].append(name)
    
    remaining_extra = extra
    
    for group_name, group_members in groups.items():
        if not group_members or remaining_extra == 0:
            continue
        
        # Sort by cumulative holidays WITH RANDOM TIE-BREAKING
        members_sorted = sorted(
            group_members,
            key=lambda n: (
                cumulative_stats.get(n, {}).get("total_holidays", 0),
                rng.random()  # Random tie-breaker!
            )
        )
        
        group_extra = min(remaining_extra, len(group_members))
        for i in range(group_extra):
            quotas[members_sorted[i]] = base + 1
            remaining_extra -= 1
    
    # Distribute remaining ONLY among people with available holidays
    if remaining_extra > 0:
        available = [n for n in available_for_holidays if quotas[n] == base]
        if available:
            available_sorted = sorted(
                available,
                key=lambda n: (
                    cumulative_stats.get(n, {}).get("total_holidays", 0),
                    rng.random()
                )
            )
            for i in range(min(remaining_extra, len(available_sorted))):
                quotas[available_sorted[i]] = base + 1
    
    return quotas


def compute_friday_quotas_with_history(
    names: list[str],
    total_fridays: int,
    cumulative_stats: dict[str, dict],
    ranks: dict[str, str],
    rng: random.Random,
    leaves: dict[str, set[int]] = None,
    year: int = None,
    month: int = None,
) -> dict[str, int]:
    """
    Compute Friday quotas with GROUP-BASED MIN-MAX balancing (±1 rule).
    Same logic as compute_holiday_quotas_with_history but for Fridays.
    - ZERO quota for people with NO available Fridays.
    """
    if leaves is None:
        leaves = {}

    # Check which people have at least one available Friday
    available_for_fridays = []
    if year and month:
        import calendar as cal
        _, days_in_month = cal.monthrange(year, month)
        for name in names:
            leave_days = leaves.get(name, set())
            has_friday = False
            for d in range(1, days_in_month + 1):
                if d not in leave_days and dt.date(year, month, d).weekday() == 4:  # Friday
                    has_friday = True
                    break
            if has_friday:
                available_for_fridays.append(name)
    else:
        available_for_fridays = names[:]

    if not available_for_fridays:
        return {name: 0 for name in names}

    base, extra = divmod(total_fridays, len(available_for_fridays))
    quotas = {name: 0 for name in names}
    for name in available_for_fridays:
        quotas[name] = base

    if extra == 0:
        return quotas

    # Group by rank for fair distribution
    groups = {"OFFICERS": [], "OTHERS": []}
    for name in available_for_fridays:
        rank = ranks.get(name, "")
        group = get_person_group(rank)
        groups[group].append(name)

    remaining_extra = extra

    for group_name, group_members in groups.items():
        if not group_members or remaining_extra == 0:
            continue

        members_sorted = sorted(
            group_members,
            key=lambda n: (
                cumulative_stats.get(n, {}).get("total_fridays", 0),
                rng.random()
            )
        )

        group_extra = min(remaining_extra, len(group_members))
        for i in range(group_extra):
            quotas[members_sorted[i]] = base + 1
            remaining_extra -= 1

    if remaining_extra > 0:
        available = [n for n in available_for_fridays if quotas[n] == base]
        if available:
            available_sorted = sorted(
                available,
                key=lambda n: (
                    cumulative_stats.get(n, {}).get("total_fridays", 0),
                    rng.random()
                )
            )
            for i in range(min(remaining_extra, len(available_sorted))):
                quotas[available_sorted[i]] = base + 1

    return quotas



# -----------------------------
# Main Solver (single tab)
# -----------------------------
# -----------------------------
# Random seed management
# -----------------------------
_solve_counter = 0  # Global counter for unique seeds

def solve_schedule_best_effort(
    names: list[str],
    year: int,
    month: int,
    extra_holidays: set[int],
    leaves: dict[str, set[int]],
    preferences: dict[str, set[int]] = None,  # soft preferences
    max_caps: dict[str, int] | None = None,
    ranks: dict[str, str] | None = None,
    tab_key: str = "AYDM",
    log_cb=None,
) -> tuple[dict[int, str], dict[str, int], dict, dict]:
    """
    Main solver with two-phase strategy and score balancing.

    MAX behavior (as requested):
    - If user sets MAX for a person (e.g. MAX=1), that person must get EXACTLY that many services.
    - The remaining days are distributed among the rest (history-aware), respecting all constraints.
    - No pre-assignment of specific days for MAX people (so we don't "lock" the solution space).
    """
    global _solve_counter
    _solve_counter += 1  # Increment for unique seed

    if log_cb is None:
        log_cb = lambda m: print(m)

    if preferences is None:
        preferences = {}

    if not names:
        raise ScheduleError("Δεν υπάρχουν άτομα!", details="Πρόσθεσε τουλάχιστον ένα άτομο.")

    if ranks is None:
        ranks = {}

    _, days_in_month = calendar.monthrange(year, month)

    # Count total holidays in month
    total_holidays = sum(
        1 for d in range(1, days_in_month + 1)
        if day_bucket(year, month, d, extra_holidays) == "HOLIDAY"
    )

    # Count total Fridays in month
    total_fridays = sum(
        1 for d in range(1, days_in_month + 1)
        if dt.date(year, month, d).weekday() == 4  # Friday
    )

    # AUTO-CALCULATE MAX for people with limited availability (caps + custom min gaps)
    log_cb("🔍 Έλεγχος διαθεσιμότητας...")

    for name in names:
        leave_days = leaves.get(name, set())
        if leave_days and DEBUG:
            available_days = set(range(1, days_in_month + 1)) - leave_days
            log_cb(
                f"  DEBUG {name}: Άδειες={sorted(leave_days)}, "
                f"Διαθέσιμες={len(available_days)} μέρες: {sorted(available_days)}"
            )

    auto_max, custom_min_gap = compute_auto_max_from_leaves(
        names, days_in_month, leaves, MIN_GAP_STRICT, log_cb
    )

    # Merge user MAX with auto MAX caps (take minimum if both exist)
    if max_caps is None:
        max_caps = {}

    combined_max_caps = max_caps.copy()
    for name, auto_val in auto_max.items():
        if name not in combined_max_caps or combined_max_caps[name] is None:
            combined_max_caps[name] = auto_val
        else:
            try:
                user_val = int(combined_max_caps[name])
                combined_max_caps[name] = min(user_val, auto_val)
            except Exception:
                combined_max_caps[name] = auto_val

    # -------- Fixed MAX (exact quota), no pre-assign --------
    fixed_max: dict[str, int] = {}
    for n, v in (max_caps or {}).items():
        if n in names and v is not None:
            try:
                fixed_max[n] = int(v)
            except Exception:
                # ignore invalid, parser should prevent this
                pass

    # Also respect auto caps for fixed max (safety)
    for n in list(fixed_max.keys()):
        if n in auto_max:
            fixed_max[n] = min(fixed_max[n], int(auto_max[n]))

    fixed_total = sum(fixed_max.values())
    if fixed_total > days_in_month:
        raise ScheduleError(
            "Αδύνατο πρόγραμμα",
            details=f"Το άθροισμα των MAX ({fixed_total}) είναι μεγαλύτερο από τις μέρες του μήνα ({days_in_month})."
        )

    fixed_people = set(fixed_max.keys())
    remaining_days = days_in_month - fixed_total
    remaining_names = [n for n in names if n not in fixed_people]

    # Helper: build total quotas with fixed MAX + history-aware distribution for the rest
    def build_quotas(rng: random.Random) -> dict[str, int]:
        quotas: dict[str, int] = {}

        # Fixed ones
        for p, q in fixed_max.items():
            quotas[p] = q

        # Distribute remaining among others
        if remaining_names:
            cumulative_stats = calculate_cumulative_stats(year, month, tab_key, remaining_names)
            if ranks and cumulative_stats:
                q_init = compute_total_quotas_with_history(
                    remaining_names, ranks, remaining_days, cumulative_stats, rng
                )
            else:
                q_init = compute_total_quotas(remaining_names, remaining_days, rng)

            q_norm = normalize_quotas_with_max(
                remaining_names, q_init, remaining_days, combined_max_caps
            )
            quotas.update(q_norm)
        else:
            # Everyone has fixed max: ensure exact fill
            if remaining_days != 0:
                raise ScheduleError(
                    "Αδύνατο πρόγραμμα",
                    details="Όλοι έχουν MAX αλλά δεν καλύπτεται ακριβώς ο μήνας. Έλεγξε τα MAX."
                )

        return quotas

    # Helper: holiday quotas that never exceed total quotas per person
    def build_holiday_quotas(
        rng: random.Random,
        quotas: dict[str, int],
        forced_holiday_counts: dict[str, int],
    ) -> dict[str, int]:
        """Holiday quotas for the *remaining* assignments.

        We compute fair holiday quotas for the whole month (history-aware),
        then subtract already-forced holiday assignments from preferences.
        """
        cumulative_stats = calculate_cumulative_stats(year, month, tab_key, names)
        holiday_quotas_total = compute_holiday_quotas_with_history(
            names, total_holidays, cumulative_stats, ranks, rng,
            leaves=leaves,
            year=year,
            month=month,
            extra_holidays=extra_holidays,
        )

        # Subtract forced holiday assignments
        holiday_quotas: dict[str, int] = {}
        for p in names:
            forced_h = int(forced_holiday_counts.get(p, 0))
            q_total = int(holiday_quotas_total.get(p, 0))
            q_remaining = max(0, q_total - forced_h)

            # Clamp by remaining total quotas (can't have more holidays than remaining assignments)
            qp = int(quotas.get(p, 0))
            q_remaining = min(q_remaining, qp)
            if qp <= 0:
                q_remaining = 0

            holiday_quotas[p] = q_remaining

        return holiday_quotas


    def build_friday_quotas(
        rng: random.Random,
        quotas: dict[str, int],
        forced_friday_counts: dict[str, int],
    ) -> dict[str, int]:
        """Friday quotas for the *remaining* assignments (±1 rule, separate from holidays)."""
        cumulative_stats = calculate_cumulative_stats(year, month, tab_key, names)
        friday_quotas_total = compute_friday_quotas_with_history(
            names, total_fridays, cumulative_stats, ranks if ranks else {},
            rng,
            leaves=leaves,
            year=year,
            month=month,
        )

        friday_quotas: dict[str, int] = {}
        for p in names:
            forced_f = int(forced_friday_counts.get(p, 0))
            q_total = int(friday_quotas_total.get(p, 0))
            q_remaining = max(0, q_total - forced_f)
            # Clamp by remaining total quotas
            qp = int(quotas.get(p, 0))
            q_remaining = min(q_remaining, qp)
            if qp <= 0:
                q_remaining = 0
            friday_quotas[p] = q_remaining

        return friday_quotas

    def build_forced_preassigned(rng: random.Random) -> dict[int, str]:
        """Turn 'preferences' into forced pre-assignments.

        Rule requested:
        - If two (or more) people request the same day, pick ONE randomly.
          The rest are ignored for that day.
        - If a preference day falls on a person's leave day, it is ignored.
        - If a person has fixed MAX and requests more than MAX days, we keep
          a random subset up to MAX (the rest ignored).
        """
        day_to_candidates: dict[int, list[str]] = {}

        for p, days in (preferences or {}).items():
            if p not in names:
                continue
            if not days:
                continue

            # Keep only valid, available days
            cand = [d for d in days if 1 <= int(d) <= days_in_month and int(d) not in leaves.get(p, set())]
            if not cand:
                continue

            # Respect fixed MAX: keep at most MAX preferred days
            lim = fixed_max.get(p)
            if lim is not None:
                try:
                    lim = int(lim)
                except Exception:
                    lim = None
            if lim is not None:
                if lim <= 0:
                    continue
                if len(cand) > lim:
                    rng.shuffle(cand)
                    cand = cand[:lim]

            for d in cand:
                day_to_candidates.setdefault(int(d), []).append(p)

        forced: dict[int, str] = {}
        for d, cands in day_to_candidates.items():
            if not cands:
                continue
            forced[int(d)] = rng.choice(cands)

        return forced
# Try multiple times with different random seeds to find best score spread
    # Phase A: target spread = 0.0 | Phase B: accept spread <= 1.0
    best_spread = float("inf")
    best_schedule = None
    best_quotas = None
    best_meta = None
    best_solve_info = {}
    last_error = None

    # Initialize fallback forced counts (will be updated by run_attempt if any attempt succeeds)
    forced_pre: dict[int, str] = {}
    forced_holiday_counts: dict[str, int] = {p: 0 for p in names}
    forced_friday_counts: dict[str, int]  = {p: 0 for p in names}

    TRIES_A = 40  # Single phase targeting spread <= 1.0

    def run_attempt(attempt_num, rng):
        nonlocal last_error, forced_pre, forced_holiday_counts, forced_friday_counts
        _forced_pre = build_forced_preassigned(rng)

        _forced_counts: dict[str, int] = {p: 0 for p in names}
        _forced_hol: dict[str, int]    = {p: 0 for p in names}
        _forced_fri: dict[str, int]    = {p: 0 for p in names}
        for d, p in _forced_pre.items():
            _forced_counts[p] = _forced_counts.get(p, 0) + 1
            if day_bucket(year, month, d, extra_holidays) == "HOLIDAY":
                _forced_hol[p] = _forced_hol.get(p, 0) + 1
            if dt.date(year, month, d).weekday() == 4:
                _forced_fri[p] = _forced_fri.get(p, 0) + 1

        _qtotal = build_quotas(rng)
        _quotas: dict[str, int] = {}
        for p in names:
            q = int(_qtotal.get(p, 0)) - int(_forced_counts.get(p, 0))
            _quotas[p] = max(0, q)

        _remaining = days_in_month - len(_forced_pre)
        _diff = _remaining - sum(_quotas.values())
        if _diff != 0:
            _adj = [p for p in names if p not in fixed_people] or names[:]
            if _diff > 0:
                for _ in range(_diff):
                    rng.shuffle(_adj)
                    _added = False
                    for p in _adj:
                        _cap = combined_max_caps.get(p)
                        if _cap is not None:
                            try: _cap = int(_cap)
                            except: _cap = None
                        _planned = int(_forced_counts.get(p,0)) + int(_quotas.get(p,0))
                        if _cap is None or _planned < _cap:
                            _quotas[p] += 1; _added = True; break
                    if not _added:
                        _quotas[rng.choice(_adj)] += 1
            else:
                _pool = [p for p in _adj if _quotas.get(p,0) > 0]
                for _ in range(-_diff):
                    if not _pool: break
                    p = rng.choice(_pool)
                    _quotas[p] -= 1
                    if _quotas[p] <= 0:
                        _pool = [x for x in _pool if x != p]

        _hol_q = build_holiday_quotas(rng, _quotas, _forced_hol)
        _fri_q = build_friday_quotas(rng, _quotas, _forced_fri)

        if attempt_num == 0:
            if DEBUG:
                log_cb(f"  DEBUG combined_max_caps: {combined_max_caps}")
                log_cb(f"  DEBUG fixed_max: {fixed_max}")
            log_cb(f"📊 Total quotas: {_quotas}")
            log_cb(f"🎉 Holiday quotas: {_hol_q} (σύνολο αργιών: {total_holidays})")
            log_cb(f"📅 Friday quotas: {_fri_q} (σύνολο παρασκευών: {total_fridays})")

        try:
            if attempt_num == 0:
                log_cb("📌 \u03A0\u03C1\u03BF\u03C3\u03C0\u03AC\u03B8\u03B5\u03B9\u03B1: Two-phase (\u0391\u03C1\u03B3\u03AF\u03B5\u03C2 \u2192 \u039A\u03B1\u03B8\u03B7\u03BC\u03B5\u03C1\u03B9\u03BD\u03AD\u03C2)...")
            _sched, _si, _meth = solve_with_two_phase_backtracking(
                names=names, year=year, month=month,
                extra_holidays=extra_holidays, leaves=leaves,
                quotas=_quotas, holiday_quotas=_hol_q, friday_quotas=_fri_q,
                preferences=preferences, min_gap=MIN_GAP_STRICT,
                custom_min_gap=custom_min_gap, pre_assigned=_forced_pre,
                log_cb=log_cb if attempt_num == 0 else (lambda m: None),
            )
            _meta = compute_schedule_metadata(_sched, year, month, extra_holidays)
            _scores = [m["score"] for m in _meta.values()]
            _spread = max(_scores) - min(_scores) if _scores else 0.0

            # Update shared forced vars so fallbacks have valid values
            forced_pre = _forced_pre
            forced_holiday_counts = _forced_hol
            forced_friday_counts  = _forced_fri
            return _spread, _sched, _quotas, _meta, _si, _forced_pre
        except ScheduleError as e:
            if attempt_num == 0:
                log_cb(f"⚠️  Two-phase αποτυχία: {e.details}")
            last_error = e
            return None

    import time as _t, os as _os

    # ── Αναζήτηση spread ≤ 1.0 ───────────────────────────────────
    log_cb(f"🔍 Αναζήτηση spread ≤ 1.0 ({TRIES_A} προσπάθειες)...")
    for _att in range(TRIES_A):
        _seed = int(_t.time()*1000000) + _att*12345 + _os.getpid() + _solve_counter*99991
        _rng = random.Random(_seed)
        _res = run_attempt(_att, _rng)
        if _res is None: continue
        _sp, _sc, _qu, _me, _si, _fp = _res
        log_cb(f"  🎯 Προσπάθεια #{_att+1}: spread = {_sp:.2f}")
        if _sp < best_spread:
            best_spread, best_schedule, best_quotas, best_meta, best_solve_info = _sp, _sc, _qu, _me, _si
        if best_spread <= 1.0:
            log_cb("  ✅ Βρέθηκε spread ≤ 1.0!")
            break

    if best_schedule is not None:
        if best_spread <= 1.0:
            log_cb(f"✅ Τελικό Score spread: {best_spread:.2f} (≤ 1.0 ✓)")
        else:
            log_cb(f"⚠️  Τελικό Score spread: {best_spread:.2f} (> 1.0, αλλά καλύτερο δυνατό)")
        return best_schedule, best_quotas, best_meta, best_solve_info


    # ------------------ Fallbacks: relax GAP only ------------------
    log_cb("\U0001F4CC Fallback 1: GAP=1 για πολύ περιορισμένους...")

    highly_constrained = {}
    for name in names:
        leave_days = leaves.get(name, set())
        available = len([d for d in range(1, days_in_month + 1) if d not in leave_days])
        if available < days_in_month / 3:
            highly_constrained[name] = 1

    if highly_constrained:
        import time
        rng = random.Random(int(time.time() * 1000) % (2**32))
        quotas = build_quotas(rng)
        holiday_quotas = build_holiday_quotas(rng, quotas, forced_holiday_counts)
        friday_quotas_fb = build_friday_quotas(rng, quotas, forced_friday_counts)
        try:
            schedule, solve_info, method = solve_with_two_phase_backtracking(
                names=names, year=year, month=month,
                extra_holidays=extra_holidays, leaves=leaves,
                quotas=quotas, holiday_quotas=holiday_quotas,
                friday_quotas=friday_quotas_fb,
                preferences=preferences, min_gap=2,
                custom_min_gap=highly_constrained, pre_assigned=forced_pre,
                log_cb=lambda m: None,
            )
            meta = compute_schedule_metadata(schedule, year, month, extra_holidays)
            log_cb(f"✅ Βρέθηκε λύση με GAP=1 για {len(highly_constrained)} άτομα")
            return schedule, quotas, meta, solve_info
        except ScheduleError as e:
            last_error = e

    log_cb("\U0001F4CC Fallback 2: GAP=1 για μέτρια περιορισμένους...")

    moderately_constrained = {}
    for name in names:
        leave_days = leaves.get(name, set())
        available = len([d for d in range(1, days_in_month + 1) if d not in leave_days])
        if available < days_in_month / 2:
            moderately_constrained[name] = 1

    if moderately_constrained:
        import time
        rng = random.Random(int(time.time() * 1000) % (2**32))
        quotas = build_quotas(rng)
        holiday_quotas = build_holiday_quotas(rng, quotas, forced_holiday_counts)
        friday_quotas_fb = build_friday_quotas(rng, quotas, forced_friday_counts)
        try:
            schedule, solve_info, method = solve_with_two_phase_backtracking(
                names=names, year=year, month=month,
                extra_holidays=extra_holidays, leaves=leaves,
                quotas=quotas, holiday_quotas=holiday_quotas,
                friday_quotas=friday_quotas_fb,
                preferences=preferences, min_gap=2,
                custom_min_gap=moderately_constrained, pre_assigned=forced_pre,
                log_cb=lambda m: None,
            )
            meta = compute_schedule_metadata(schedule, year, month, extra_holidays)
            log_cb(f"✅ Βρέθηκε λύση με GAP=1 για {len(moderately_constrained)} άτομα")
            return schedule, quotas, meta, solve_info
        except ScheduleError as e:
            last_error = e

    log_cb("\U0001F4CC Fallback 3: GAP=1 για όλους...")

    import time
    rng = random.Random(int(time.time() * 1000) % (2**32))
    quotas = build_quotas(rng)
    holiday_quotas = build_holiday_quotas(rng, quotas, forced_holiday_counts)
    friday_quotas_fb = build_friday_quotas(rng, quotas, forced_friday_counts)
    try:
        schedule, solve_info, method = solve_with_two_phase_backtracking(
            names=names, year=year, month=month,
            extra_holidays=extra_holidays, leaves=leaves,
            quotas=quotas, holiday_quotas=holiday_quotas,
            friday_quotas=friday_quotas_fb,
            preferences=preferences, min_gap=1,
            custom_min_gap=None, pre_assigned=forced_pre,
            log_cb=lambda m: None,
        )
        meta = compute_schedule_metadata(schedule, year, month, extra_holidays)
        log_cb("✅ Βρέθηκε λύση με GAP=1 για \u03CC\u03BB\u03BF\u03C5\u03C2")
        return schedule, quotas, meta, solve_info
    except ScheduleError as e:
        last_error = e


    # ── PROGRESSIVE CONSTRAINT RELAXATION ────────────────────────
    # Step 1: Relax Friday quota constraint (try completely without Friday balancing)

    log_cb("🔧 Φάση Γ: Χαλάρωση κανόνα Παρασκευών...")

    def try_without_friday_quota():
        import time as _tfy
        _best_sp = float("inf")
        _best_res = None
        for _xi in range(20):
            _xseed = int(_tfy.time()*1000000) + _xi*44444
            _xrng  = random.Random(_xseed)
            _xq    = build_quotas(_xrng)
            _xhq   = build_holiday_quotas(_xrng, _xq, forced_holiday_counts)
            # No friday quota - pass all 999
            _xfq   = {p: 999 for p in names}
            try:
                _sc, _si, _ = solve_with_two_phase_backtracking(
                    names=names, year=year, month=month, extra_holidays=extra_holidays,
                    leaves=leaves, quotas=_xq, holiday_quotas=_xhq, friday_quotas=_xfq,
                    preferences=preferences, min_gap=MIN_GAP_STRICT,
                    custom_min_gap=custom_min_gap, pre_assigned=forced_pre,
                    log_cb=lambda m: None,
                )
                _meta = compute_schedule_metadata(_sc, year, month, extra_holidays)
                _scores = [m["score"] for m in _meta.values()]
                _sp = max(_scores) - min(_scores) if _scores else 0.0
                if _sp < _best_sp:
                    _best_sp = _sp
                    _best_res = (_sc, _xq, _meta, _si)
                if _best_sp <= 1.0:
                    break
            except ScheduleError:
                continue
        return _best_res, _best_sp

    _res_fri, _sp_fri = try_without_friday_quota()
    if _res_fri is not None:
        _sc_fri, _q_fri, _m_fri, _si_fri = _res_fri
        log_cb(f"  ⚠️  ΠΡΟΣΟΧΗ: Κανόνας Παρασκευών αγνοήθηκε")
        log_cb(f"✅ Βρέθηκε πρόγραμμα χωρίς Friday quota (spread={_sp_fri:.2f})")
        _si_fri['relaxed_friday_quota'] = True
        return _sc_fri, _q_fri, _m_fri, _si_fri

    # Step 2: Relax Sat/Fri pair rule only for the specific days that need it

    log_cb("🔧 Φάση Δ: Χαλάρωση κανόνα Παρ/Σαβ για συγκεκριμένες μέρες...")

    def try_with_relaxed_weekend_pair(ignore_days):
        import time as _tw
        _best_sp = float("inf")
        _best_res = None
        for _xi in range(20):
            _xseed = int(_tw.time()*1000000) + _xi*33333
            _xrng  = random.Random(_xseed)
            _xq    = build_quotas(_xrng)
            _xhq   = build_holiday_quotas(_xrng, _xq, forced_holiday_counts)
            _xfq   = build_friday_quotas(_xrng, _xq, forced_friday_counts)
            try:
                _sc, _si, _ = solve_with_two_phase_backtracking(
                    names=names, year=year, month=month, extra_holidays=extra_holidays,
                    leaves=leaves, quotas=_xq, holiday_quotas=_xhq, friday_quotas=_xfq,
                    preferences=preferences, min_gap=MIN_GAP_STRICT,
                    custom_min_gap=custom_min_gap, pre_assigned=forced_pre,
                    log_cb=lambda m: None,
                    ignore_weekend_pair_days=ignore_days,
                )
                _meta = compute_schedule_metadata(_sc, year, month, extra_holidays)
                _scores = [m["score"] for m in _meta.values()]
                _sp = max(_scores) - min(_scores) if _scores else 0.0
                if _sp < _best_sp:
                    _best_sp = _sp
                    _best_res = (_sc, _xq, _meta, _si)
                if _best_sp <= 1.0:
                    break
            except ScheduleError:
                continue
        return _best_res, _best_sp

    _, _dim_r = calendar.monthrange(year, month)
    _friday_days_r = {d for d in range(1, _dim_r+1) if dt.date(year, month, d).weekday() == 4}

    # Find days mentioned in last_error with "Κανόνας Σαβ/Παρ"
    _weekend_problem_days = set()
    if last_error and hasattr(last_error, 'details') and last_error.details:
        import re
        for _match in re.finditer(r'(\d+)/' + str(month), last_error.details):
            _d = int(_match.group(1))
            # Check if this day has a Σαβ/Παρ mention nearby
            _pos = _match.start()
            _context = last_error.details[max(0,_pos-5):_pos+50]
            if 'Κανόνας Σαβ/Παρ' in last_error.details[_pos:_pos+200]:
                _weekend_problem_days.add(_d)
        # Also add: days that appear in "few_cand" section with Σαβ/Παρ exclusions
        for _d in range(1, _dim_r+1):
            _day_str = f"{_d}/{month}"
            _idx = last_error.details.find(_day_str)
            if _idx >= 0:
                _section = last_error.details[_idx:_idx+400]
                if 'Κανόνας Σαβ/Παρ' in _section:
                    _weekend_problem_days.add(_d)

    if not _weekend_problem_days:
        # Fallback: try all Fridays
        _weekend_problem_days = _friday_days_r

    log_cb(f"  🔧 Χαλάρωση Σαβ/Παρ για μέρες: {sorted(_weekend_problem_days)}")
    _res_r, _sp_r = try_with_relaxed_weekend_pair(_weekend_problem_days)

    if _res_r is not None:
        _sc_r, _q_r, _m_r, _si_r = _res_r
        _relaxed_str = ", ".join(f"{d}/{month}" for d in sorted(_weekend_problem_days))
        log_cb(f"  ⚠️  ΠΡΟΣΟΧΗ: Κανόνας Σαβ/Παρ αγνοήθηκε για: {_relaxed_str}")
        log_cb(f"✅ Βρέθηκε πρόγραμμα με χαλαρωμένο Σαβ/Παρ (spread={_sp_r:.2f})")
        _si_r['relaxed_weekend_pair_days'] = sorted(_weekend_problem_days)
        _si_r['relaxed_friday_quota'] = not enforce_friday_quota if 'enforce_friday_quota' in dir() else False
        return _sc_r, _q_r, _m_r, _si_r

    # ── FEASIBILITY PROOF ─────────────────────────────────────────
    # Prove it is truly impossible before raising error.
    # Use min_gap=1, no friday cap, extended timeout.
    log_cb("\U0001F52C Έλεγχος αδυναμίας (min_gap=1, χαλαρά όρια)...")

    import sys as _sys
    _mod = _sys.modules[__name__]
    _orig_time = _mod.SOLVER_MAX_TIME
    _orig_mult = _mod.SOLVER_MAX_RECURSION_MULT
    _mod.SOLVER_MAX_TIME = 60.0
    _mod.SOLVER_MAX_RECURSION_MULT = 10

    _proof_rng = random.Random(42)
    _proof_q   = build_quotas(_proof_rng)
    _proof_hq  = build_holiday_quotas(_proof_rng, _proof_q, forced_holiday_counts)
    _feasible  = False
    try:
        solve_with_two_phase_backtracking(
            names=names, year=year, month=month,
            extra_holidays=extra_holidays, leaves=leaves,
            quotas=_proof_q, holiday_quotas=_proof_hq,
            friday_quotas={p: 999 for p in names},
            preferences=None, min_gap=1,
            custom_min_gap=None, pre_assigned={},
            log_cb=lambda m: None,
        )
        _feasible = True
    except ScheduleError:
        _feasible = False
    finally:
        _mod.SOLVER_MAX_TIME = _orig_time
        _mod.SOLVER_MAX_RECURSION_MULT = _orig_mult

    if _feasible:
        # Solvable but needs more attempts
        log_cb("\u26A0\uFE0F  \u03A4\u03BF \u03C0\u03C1\u03CC\u03B3\u03C1\u03B1\u03BC\u03BC\u03B1 \u0392\u0393\u0391\u0399\u039D\u0395\u0399 \u03B1\u03BB\u03BB\u03AC \u03C7\u03C1\u03B5\u03B9\u03AC\u03B6\u03B5\u03C4\u03B1\u03B9 \u03C0\u03B5\u03C1\u03B9\u03C3\u03C3\u03CC\u03C4\u03B5\u03C1\u03B5\u03C2 προσπάθειες. \u0395\u03C0\u03B1\u03BD\u03B1\u03C0\u03C1\u03BF\u03C3\u03C0\u03AC\u03B8\u03B5\u03B9\u03B1...")
        _mod.SOLVER_MAX_TIME = 60.0
        _mod.SOLVER_MAX_RECURSION_MULT = 8
        import time as _t2, os as _os2
        _xbest = None
        _xspread = float("inf")
        for _xi in range(50):
            _xseed = int(_t2.time()*1000000) + _xi*77777 + _solve_counter*11111
            _xrng = random.Random(_xseed)
            _xres = run_attempt(200+_xi, _xrng)
            if _xres is None: continue
            _xsp, _xsc, _xqu, _xme, _xsi, _xfp = _xres
            log_cb(f"  \U0001F504 Επαναπροσπάθεια #{_xi+1}: spread = {_xsp:.2f}")
            if _xsp < _xspread:
                _xspread = _xsp; _xbest = (_xsc, _xqu, _xme, _xsi)
            if _xspread <= 1.0: break
        _mod.SOLVER_MAX_TIME = _orig_time
        _mod.SOLVER_MAX_RECURSION_MULT = _orig_mult
        if _xbest is not None:
            log_cb(f"✅ Βρέθηκε λύση με spread = {_xspread:.2f}")
            return _xbest[0], _xbest[1], _xbest[2], _xbest[3]

    # Truly impossible
    if last_error is not None:
        raise last_error
    raise ScheduleError(
        "Αδύνατο πρόγραμμα",
        details="Δεν βρέθηκε λύση. Έλεγξε άδειες και constraints."
    )
def compute_schedule_metadata(
    schedule: dict[int, str],
    year: int,
    month: int,
    extra_holidays: set[int],
) -> dict:
    """Compute stats per person with actual dates."""
    meta = {}
    
    for day, person in schedule.items():
        if person not in meta:
            meta[person] = {
                "WEEKDAY": 0,  # Mon-Thu only
                "FRIDAY": 0,   # Friday separately
                "HOLIDAY": 0,  # Sat/Sun/Holidays
                "total": 0, 
                "score": 0.0,
                "weekday_dates": [],  # Mon-Thu dates
                "friday_dates": [],   # Friday dates
                "holiday_dates": [],  # Holiday dates
            }
        
        # Determine day type
        date_obj = dt.date(year, month, day)
        weekday_num = date_obj.weekday()  # 0=Mon, 4=Fri, 5=Sat, 6=Sun
        bucket = day_bucket(year, month, day, extra_holidays)
        
        if bucket == "HOLIDAY":
            meta[person]["HOLIDAY"] += 1
            meta[person]["holiday_dates"].append(day)
            meta[person]["score"] += H_SCORE
        elif weekday_num == 4:  # Friday
            meta[person]["FRIDAY"] += 1
            meta[person]["friday_dates"].append(day)
            meta[person]["score"] += W_SCORE
        else:  # Mon-Thu
            meta[person]["WEEKDAY"] += 1
            meta[person]["weekday_dates"].append(day)
            meta[person]["score"] += W_SCORE
        
        meta[person]["total"] += 1
    
    return meta


# -----------------------------
# Export to Word
# -----------------------------
def set_cell_background(cell, color_rgb: tuple[int, int, int]):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), f'{color_rgb[0]:02X}{color_rgb[1]:02X}{color_rgb[2]:02X}')
    cell._element.get_or_add_tcPr().append(shading)



def export_all_schedules_to_word(
    filepath: str,
    export_data: list[dict],
    year: int,
    month: int,
    extra_holidays: set[int],
):
    """
    Export all schedules to a single Word file with page breaks.
    """
    from docx.oxml import OxmlElement as _OE

    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2)
        section.right_margin  = Cm(2)

    def set_font(run, size_pt, bold=False):
        run.font.name = "Cambria"
        run.font.size = Pt(size_pt)
        run.bold = bold

    def set_row_height(row, h_cm=0.5):
        trPr = row._tr.get_or_add_trPr()
        trH = _OE("w:trHeight")
        trH.set(qn("w:val"), str(int(Cm(h_cm).twips)))
        trH.set(qn("w:hRule"), "exact")
        trPr.append(trH)

    def set_cell_width(cell, w_cm):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = _OE("w:tcW")
        tcW.set(qn("w:w"), str(int(Cm(w_cm).twips)))
        tcW.set(qn("w:type"), "dxa")
        existing = tcPr.find(qn("w:tcW"))
        if existing is not None:
            tcPr.remove(existing)
        tcPr.insert(0, tcW)

    def set_cell_valign_center(cell):
        tcPr = cell._tc.get_or_add_tcPr()
        vAlign = _OE("w:vAlign")
        vAlign.set(qn("w:val"), "center")
        tcPr.append(vAlign)

    col_widths = [3, 3, 3, 5.5]
    _, days_in_month = calendar.monthrange(year, month)

    for idx, data in enumerate(export_data):
        tab_key = data["tab_key"]
        schedule = data["schedule"]
        ranks    = data["ranks"]

        if idx > 0:
            doc.add_page_break()

        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = title.add_run("ΠΙΝΑΚΑΣ")
        set_font(r, 11, bold=True)
        r.underline = True

        # Subtitle
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        month_name_upper = GREEK_MONTHS_GEN[month].upper()
        r = subtitle.add_run(f"ΥΠΗΡΕΣΙΩΝ {TAB_TITLES.get(tab_key, tab_key)} ΜΗΝΟΣ {month_name_upper} {year}")
        set_font(r, 11, bold=True)

        doc.add_paragraph()  # blank line

        # Table
        table = doc.add_table(rows=days_in_month + 1, cols=4)
        table.style = "Table Grid"

        # Column widths via tblGrid
        tbl = table._tbl
        tblGrid = _OE("w:tblGrid")
        for w_cm in col_widths:
            gc = _OE("w:gridCol")
            gc.set(qn("w:w"), str(int(Cm(w_cm).twips)))
            tblGrid.append(gc)
        existing_grid = tbl.find(qn("w:tblGrid"))
        if existing_grid is not None:
            tbl.remove(existing_grid)
        tbl.insert(1, tblGrid)

        # Header row
        hdr_row = table.rows[0]
        set_row_height(hdr_row)
        hdr_labels = ["ΗΜΕΡΟΜΗΝΙΑ", "ΗΜΕΡΑ", "ΒΑΘΜΟΣ", "ΟΝΟΜΑΤΕΠΩΝΥΜΟ"]
        for ci, cell in enumerate(hdr_row.cells):
            set_cell_width(cell, col_widths[ci])
            set_cell_valign_center(cell)
            cell.text = ""
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = para.add_run(hdr_labels[ci])
            set_font(r, 11, bold=True)
            shading = parse_xml(r'<w:shd {} w:fill="D3D3D3"/>'.format(nsdecls("w")))
            cell._element.get_or_add_tcPr().append(shading)

        # Data rows
        for day in range(1, days_in_month + 1):
            row = table.rows[day]
            set_row_height(row)
            cells = row.cells

            date_obj   = dt.date(year, month, day)
            weekday_gr = {
                "Monday": "ΔΕΥΤΕΡΑ", "Tuesday": "ΤΡΙΤΗ", "Wednesday": "ΤΕΤΑΡΤΗ",
                "Thursday": "ΠΕΜΠΤΗ", "Friday": "ΠΑΡΑΣΚΕΥΗ",
                "Saturday": "ΣΑΒΒΑΤΟ", "Sunday": "ΚΥΡΙΑΚΗ"
            }.get(date_obj.strftime("%A"), "")

            person   = schedule.get(day, "???")
            rank     = ranks.get(person, "")
            date_str = f"{day}-{GREEK_MONTH_ABBR[month].upper()}-{str(year)[-2:]}"
            row_vals = [date_str, weekday_gr, rank, person.upper()]

            for ci, cell in enumerate(cells):
                set_cell_width(cell, col_widths[ci])
                set_cell_valign_center(cell)
                cell.text = ""
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = para.add_run(row_vals[ci])
                set_font(r, 11)

            # Gray background for holidays
            if day_bucket(year, month, day, extra_holidays) == "HOLIDAY":
                for cell in cells:
                    shading = parse_xml(r'<w:shd {} w:fill="D3D3D3"/>'.format(nsdecls("w")))
                    cell._element.get_or_add_tcPr().append(shading)

    doc.save(filepath)


# -----------------------------
# GUI: PeopleTab
# -----------------------------
class PeopleTab(ttk.Frame):
    def __init__(self, parent, tab_key: str, app_ref):
        super().__init__(parent)
        self.tab_key = tab_key
        self.app_ref = app_ref
        
        self.current_schedule: dict[int, str] = {}
        self.current_ranks: dict[str, str] = {}
        
        self._build_ui()
    
    def _build_ui(self):
        top = ttk.Frame(self)
        top.pack(fill="both", expand=True, padx=10, pady=10)
        
        left = ttk.LabelFrame(top, text="Προσωπικό", padding=10)
        left.pack(side="left", fill="both", expand=True, padx=(0, 5))
        
        count_frame = ttk.Frame(left)
        count_frame.pack(fill="x", pady=(0, 10))
        ttk.Label(count_frame, text="Πλήθος ατόμων:").pack(side="left")
        
        self.count_var = tk.StringVar(value="1")
        self.active_count = tk.IntVar(value=1)
        
        count_entry = ttk.Entry(count_frame, textvariable=self.count_var, width=6)
        count_entry.pack(side="left", padx=6)
        
        ttk.Button(count_frame, text="Εφαρμογή", command=self.apply_count).pack(side="left")
        ttk.Button(count_frame, text="💾 Αποθήκευση State", command=lambda: self.app_ref.on_manual_save_state()).pack(side="left", padx=(12, 0))
        
        canvas_frame = ttk.Frame(left)
        canvas_frame.pack(fill="both", expand=True)
        
        canvas = tk.Canvas(canvas_frame, height=300)
        scrollbar = ttk.Scrollbar(canvas_frame, orient="vertical", command=canvas.yview)
        self.scrollable = ttk.Frame(canvas)
        
        self.scrollable.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=self.scrollable, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        hdr = ttk.Frame(self.scrollable)
        hdr.pack(fill="x", pady=(0, 5))
        ttk.Label(hdr, text="Βαθμός", width=12).grid(row=0, column=0, padx=2)
        ttk.Label(hdr, text="Ονοματεπώνυμο", width=20).grid(row=0, column=1, padx=2)
        ttk.Label(hdr, text="Άδειες (π.χ. 6,25)", width=18).grid(row=0, column=2, padx=2)
        
        # Button to clear all leaves
        clear_leaves_btn = ttk.Button(hdr, text="🗑️", width=3, 
                                     command=self.clear_all_leaves)
        clear_leaves_btn.grid(row=0, column=3, padx=2)
        
        ttk.Label(hdr, text="Επιθυμία (π.χ. 5,12)", width=18).grid(row=0, column=4, padx=2)
        ttk.Label(hdr, text="MAX", width=6).grid(row=0, column=5, padx=2)
        
        self.people_rows: list[dict] = []
        self._add_row("", "", "", "")
        
        # Right panel: Program / Stats
        right = ttk.LabelFrame(top, text="Πρόγραμμα / Στατιστικά", padding=10)
        right.pack(side="right", fill="both", expand=True, padx=(5, 0))

        # Toolbar (buttons live inside the Program/Stats panel)
        toolbar = ttk.Frame(right)
        toolbar.pack(fill="x", pady=(0, 8))

        ttk.Button(
            toolbar,
            text="Δημιουργία",
            command=lambda: self.app_ref.on_generate_current(),
        ).pack(side="left", padx=(0, 6))

        self.export_btn = ttk.Button(
            toolbar,
            text="💾 Αποθήκευση",
            command=lambda: self.app_ref.on_save_current_schedule(),
            state="disabled",
        )
        self.export_btn.pack(side="left")
        
        self.stats = tk.Text(right, height=35, width=60, wrap="word", state="disabled")
        self.stats.tag_configure("warning", foreground="darkorange", font=("Consolas", 9, "bold"))
        self.stats.pack(fill="both", expand=True)
        
        self.stats.tag_config("warn_total", background="yellow")
        self.stats.tag_config("info", foreground="blue")
        self.stats.tag_config("error", foreground="red")
    
    def clear_all_leaves(self):
        """Clear all leave entries."""
        for row in self.people_rows:
            row["leave_var"].set("")
    
    def _add_row(self, rank: str, name: str, leave: str, max_val: str, preference: str = ""):
        row_frame = ttk.Frame(self.scrollable)
        row_frame.pack(fill="x", pady=2)
        
        rank_var = tk.StringVar(value=rank)
        name_var = tk.StringVar(value=name)
        max_var = tk.StringVar(value=max_val)
        leave_var = tk.StringVar(value=leave)
        preference_var = tk.StringVar(value=preference)
        
        ttk.Entry(row_frame, textvariable=rank_var, width=12).grid(row=0, column=0, padx=2)
        ttk.Entry(row_frame, textvariable=name_var, width=20).grid(row=0, column=1, padx=2)
        ttk.Entry(row_frame, textvariable=leave_var, width=18).grid(row=0, column=2, padx=2)
        
        # Empty space for the clear button (column 3)
        ttk.Label(row_frame, text="", width=3).grid(row=0, column=3, padx=2)
        
        ttk.Entry(row_frame, textvariable=preference_var, width=18).grid(row=0, column=4, padx=2)
        ttk.Entry(row_frame, textvariable=max_var, width=6).grid(row=0, column=5, padx=2)
        
        self.people_rows.append({
            "frame": row_frame,
            "rank_var": rank_var,
            "name_var": name_var,
            "max_var": max_var,
            "leave_var": leave_var,
            "preference_var": preference_var,
        })

    def apply_count(self):
        """Apply the desired number of visible rows.

        - Increasing: adds empty rows.
        - Decreasing: removes ONLY empty rows (anywhere in the list). If there aren't
          enough empty rows to reach the requested count, we reject the change.
        """
        try:
            desired = safe_int(self.count_var.get(), 1)
            if desired < 1:
                desired = 1

            current = len(self.people_rows)

            # If decreasing, remove empty rows only (any position)
            if desired < current:
                need_remove = current - desired

                def is_row_empty(row: dict) -> bool:
                    return (
                        not row["name_var"].get().strip()
                        and not row["rank_var"].get().strip()
                        and not row["leave_var"].get().strip()
                        and not row["preference_var"].get().strip()
                        and not row["max_var"].get().strip()
                    )

                empty_indices = [i for i, r in enumerate(self.people_rows) if is_row_empty(r)]

                if len(empty_indices) < need_remove:
                    # Reject: user is trying to remove non-empty rows
                    self.count_var.set(str(current))
                    self.active_count.set(current)
                    messagebox.showwarning(
                        "Προσοχή",
                        "Δεν μπορώ να μειώσω το πλήθος γιατί δεν υπάρχουν αρκετές κενές γραμμές."
                        "Άδειασε πρώτα όποιες γραμμές δεν χρειάζεσαι και ξαναδοκίμασε."
                    )
                    return

                # Remove empty rows starting from the bottom-most indices to keep indices stable
                for idx in sorted(empty_indices, reverse=True)[:need_remove]:
                    try:
                        self.people_rows[idx]["frame"].destroy()
                    except Exception:
                        pass
                    del self.people_rows[idx]

                self.active_count.set(desired)
                self.count_var.set(str(desired))
                return

            # If increasing, add rows
            if desired > current:
                while len(self.people_rows) < desired:
                    self._add_row("", "", "", "", "")
                self.active_count.set(desired)
                self.count_var.set(str(desired))
                return

            # Same number: just sync vars
            self.active_count.set(current)
            self.count_var.set(str(current))

        except Exception as e:
            messagebox.showerror("Σφάλμα", str(e))

    
    
    def parse_people(self, days_in_month: int) -> tuple[list[str], dict[str, set[int]], dict[str, str], dict[str, int], dict[str, set[int]]]:
        n = self.active_count.get()
        names = []
        leaves = {}
        ranks = {}
        max_caps = {}
        preferences = {}
        
        for i in range(min(n, len(self.people_rows))):
            row = self.people_rows[i]
            name = row["name_var"].get().strip()
            if not name:
                raise ValueError(f"Το άτομο #{i+1} δεν έχει όνομα.")
            
            names.append(name)
            ranks[name] = row["rank_var"].get().strip()
            
            leave_str = row["leave_var"].get().strip()
            if leave_str:
                leaves[name] = parse_days_list(leave_str, days_in_month)
            
            preference_str = row["preference_var"].get().strip()
            if preference_str:
                preferences[name] = parse_days_list(preference_str, days_in_month)
            
            max_str = row["max_var"].get().strip()
            if max_str:
                try:
                    max_caps[name] = int(max_str)
                except ValueError:
                    raise ValueError(f"Το MAX για {name} δεν είναι αριθμός: {max_str}")
        
        return names, leaves, ranks, max_caps, preferences
    
    def collect_data(self):
        """Alias for parse_people - used by unified scheduling."""
        try:
            year = safe_int(self.app_ref.year_var.get(), 2026)
            month = safe_int(self.app_ref.month_var.get(), 3)
            _, days_in_month = calendar.monthrange(year, month)
            return self.parse_people(days_in_month)
        except Exception as e:
            raise ValueError(f"Σφάλμα στο {TAB_TITLES.get(self.tab_key, self.tab_key)}: {str(e)}")
    
    def clear_program_view(self):
        self.stats.config(state="normal")
        self.stats.delete("1.0", "end")
        self.stats.config(state="disabled")
        self.current_schedule = {}
        self.current_ranks = {}
        self.export_btn.config(state="disabled")
    
    def log(self, message: str, kind: str = "info"):
        self.stats.config(state="normal")
        start = self.stats.index("end-1c")
        self.stats.insert("end", message + "\n")
        end = self.stats.index("end-1c")
        
        if kind == "error":
            self.stats.tag_add("error", start, end)
        elif kind == "info":
            self.stats.tag_add("info", start, end)
        
        self.stats.see("end")
        self.stats.config(state="disabled")
        self.stats.update()
    
    def render_program(self, year, month, extra, sched, ranks, quotas, meta, solve_info):
        self.current_schedule = sched
        self.current_ranks = ranks
        
        self.stats.config(state="normal")
        self.stats.delete("1.0", "end")
        
        self.stats.insert("end", f"🔧 Solver: {solve_info.get('solver', 'Unknown')}\n")
        if "time" in solve_info:
            self.stats.insert("end", f"⏱️  Χρόνος: {solve_info['time']}\n")
        if "status" in solve_info:
            self.stats.insert("end", f"📊 Status: {solve_info['status']}\n")

        # Show constraint relaxation warnings prominently
        if solve_info.get('relaxed_friday_quota'):
            self.stats.insert("end", "⚠️  ΠΡΟΣΟΧΗ: Κανόνας ισοκατανομής Παρασκευών ΑΓΝΟΗΘΗΚΕ\n", "warning")
            self.stats.insert("end", "    (Το πρόγραμμα δεν έβγαινε με ισοκατανομή Παρασκευών)\n", "warning")
        if solve_info.get('relaxed_weekend_pair_days'):
            _rdays = solve_info['relaxed_weekend_pair_days']
            _rstr  = ", ".join(f"{d}/{month}" for d in _rdays)
            self.stats.insert("end", f"⚠️  ΠΡΟΣΟΧΗ: Κανόνας Σαβ/Παρ ΑΓΝΟΗΘΗΚΕ για: {_rstr}\n", "warning")
            self.stats.insert("end", "    (Το πρόγραμμα δεν έβγαινε χωρίς αυτή τη χαλάρωση)\n", "warning")

        self.stats.insert("end", "\n")
        
        _, days_in_month = calendar.monthrange(year, month)
        for day in range(1, days_in_month + 1):
            person = sched.get(day, "???")
            rank = ranks.get(person, "")
            
            # Show actual weekday instead of ΑΡΓΙΑ/ΚΑΘΗΜΕΡΙΝΗ
            date_obj = dt.date(year, month, day)
            weekday = date_obj.strftime("%A")
            weekday_gr = {
                "Monday": "ΔΕΥΤΕΡΑ", "Tuesday": "ΤΡΙΤΗ", "Wednesday": "ΤΕΤΑΡΤΗ",
                "Thursday": "ΠΕΜΠΤΗ", "Friday": "ΠΑΡΑΣΚΕΥΗ",
                "Saturday": "ΣΑΒΒΑΤΟ", "Sunday": "ΚΥΡΙΑΚΗ"
            }.get(weekday, weekday)
            
            date_str = f"{day:2d}-{GREEK_MONTH_ABBR[month]}"
            line = f"{date_str} ({weekday_gr:12s}): {rank} {person}\n"
            self.stats.insert("end", line)
        
        self.stats.insert("end", "\n" + "="*105 + "\n")
        self.stats.insert("end", "ΣΤΑΤΙΣΤΙΚΑ\n")
        self.stats.insert("end", "="*105 + "\n\n")
        
        # Table header with proper alignment
        header = f"{'Βαθμός':<15} {'Όνομα':<20} {'Καθημ.':<10} {'Παρ.':<10} {'Αργίες':<10} {'TOTAL':<10} {'POINTS':<10}\n"
        self.stats.insert("end", header)
        self.stats.insert("end", "-"*105 + "\n")
        
        scores = []
        for person in sorted(meta.keys()):
            m = meta[person]
            total = m["total"]
            weekday = m["WEEKDAY"]  # Mon-Thu
            friday = m["FRIDAY"]    # Friday
            holiday = m["HOLIDAY"]  # Sat/Sun/Holidays
            score = m["score"]
            scores.append(score)
            
            rank = ranks.get(person, "")
            
            # Format as table row with aligned columns
            row = f"{rank:<15} {person:<20} {weekday:<10} {friday:<10} {holiday:<10} {total:<10} {score:<10.1f}\n"
            
            start = self.stats.index("end-1c")
            self.stats.insert("end", row)
            end = self.stats.index("end-1c")
            
            if total >= 5:
                self.stats.tag_add("warn_total", start, end)
        
        if scores:
            self.stats.insert("end", f"\nScore spread (max-min) = {max(scores) - min(scores):.2f}\n")
            # Spread only for people with 3+ total services
            _scores_3plus = [meta[p]["score"] for p in meta if meta[p]["total"] >= 3]
            if _scores_3plus and len(_scores_3plus) < len(scores):
                self.stats.insert("end", f"Score spread (≥3 υπηρεσίες) = {max(_scores_3plus) - min(_scores_3plus):.2f}\n")
        
        # Summary of MIN_GAP usage
        gap_info = solve_info.get('min_gap_used', {})
        if gap_info:
            gap_1_users = [name for name, gap in gap_info.items() if gap == 1]
            gap_2_users = [name for name, gap in gap_info.items() if gap == 2]
            
            self.stats.insert("end", "\n" + "-"*50 + "\n")
            self.stats.insert("end", "MIN_GAP ΧΡΗΣΗ\n")
            self.stats.insert("end", "-"*50 + "\n")
            
            if gap_2_users:
                self.stats.insert("end", f"GAP=2 (κανονικό): {', '.join(gap_2_users)}\n")
            
            if gap_1_users:
                self.stats.insert("end", f"GAP=1 (χαλαρό): {', '.join(gap_1_users)}\n")
                self.stats.insert("end", f"  ℹ️  {len(gap_1_users)} άτομ{'ο' if len(gap_1_users) == 1 else 'α'} με περιορισμένη διαθεσιμότητα\n")
        
        self.stats.config(state="disabled")
        self.export_btn.config(state="normal")
    
    def set_tab_state(self, tab_state: dict):
        count = safe_int(str(tab_state.get("count", 0)), 0)
        people = tab_state.get("people", [])
        if not isinstance(people, list):
            people = []
        
        while len(self.people_rows) < max(1, count):
            self._add_row("", "", "", "", "")
        
        for i in range(max(0, count)):
            if i < len(people):
                p = people[i] if isinstance(people[i], dict) else {}
                self.people_rows[i]["rank_var"].set(str(p.get("rank", "")).strip())
                self.people_rows[i]["name_var"].set(str(p.get("name", "")).strip())
                self.people_rows[i]["max_var"].set(str(p.get("max", "")).strip())
                self.people_rows[i]["leave_var"].set(str(p.get("leave", "")).strip())
                self.people_rows[i]["preference_var"].set(str(p.get("preference", "")).strip())
            else:
                self.people_rows[i]["rank_var"].set("")
                self.people_rows[i]["name_var"].set("")
                self.people_rows[i]["max_var"].set("")
                self.people_rows[i]["leave_var"].set("")
                self.people_rows[i]["preference_var"].set("")
        
        if count <= 0:
            count = 1
        self.count_var.set(str(count))
        self.active_count.set(count)
        self.apply_count()
        
        self.clear_program_view()
        
        # Restore saved schedule/ranks/display text if present
        if "saved_schedule" in tab_state and tab_state["saved_schedule"]:
            self.current_schedule = {int(k): v for k, v in tab_state["saved_schedule"].items()}
            self.current_ranks = tab_state.get("saved_ranks", {})
            saved_text = tab_state.get("saved_stats_text", "")
            if saved_text:
                try:
                    self.stats.config(state="normal")
                    self.stats.delete("1.0", "end")
                    self.stats.insert("end", saved_text)
                    self.stats.config(state="disabled")
                    self.export_btn.config(state="normal")
                except Exception:
                    pass
    
    

    def get_tab_state(self) -> dict:
        n = self.active_count.get()
        if n < 1:
            n = 1
        self.count_var.set(str(n))
        people = []
        for i in range(min(n, len(self.people_rows))):
            row = self.people_rows[i]
            people.append({
                "rank": row["rank_var"].get().strip(),
                "name": row["name_var"].get().strip(),
                "max": row["max_var"].get().strip(),
                "leave": row["leave_var"].get().strip(),
                "preference": row["preference_var"].get().strip(),
            })
        state = {"count": n, "people": people}
        # Also save the current schedule/ranks/display text if exists
        if self.current_schedule:
            state["saved_schedule"] = {str(k): v for k, v in self.current_schedule.items()}
            state["saved_ranks"] = self.current_ranks
            # Save the stats text for display on reload
            try:
                self.stats.config(state="normal")
                state["saved_stats_text"] = self.stats.get("1.0", "end")
                self.stats.config(state="disabled")
            except Exception:
                pass
        return state


# -----------------------------
# History Viewer + Editor Dialog (AGGREGATED VIEW)
# -----------------------------
class SchedulerApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Scheduler - Two-Phase Backtracking")
        self.geometry("1250x820")
        
        self.year_var = tk.StringVar(value="2026")
        self.month_var = tk.StringVar(value="3")
        self.extra_var = tk.StringVar(value="")
        
        self._build_menu()
        self._build_ui()
        self.load_state()
        
        self.protocol("WM_DELETE_WINDOW", self.on_close)
    
    def _build_menu(self):
        """Build menu bar with History options."""
        menubar = tk.Menu(self)
        self.config(menu=menubar)
        
        # History menu
        history_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="History", menu=history_menu)
        
        history_menu.add_command(label="📝 Προβολή & Επεξεργασία...", command=self.on_edit_history)
        history_menu.add_separator()
        history_menu.add_command(label="🔄 Ανανέωση από αρχείο", command=self.on_reload_history)
        history_menu.add_command(label="📂 Άνοιγμα αρχείου JSON", command=self.on_open_history_file)
        history_menu.add_separator()
        history_menu.add_command(label="🗑️ Καθαρισμός Ιστορικού...", command=self.on_clear_history)
    
    def _build_ui(self):
        root = ttk.Frame(self, padding=10)
        root.pack(fill="both", expand=True)
        
        top = ttk.LabelFrame(root, text="Κοινές Ρυθμίσεις (για ΟΛΑ τα φύλλα)", padding=10)
        top.pack(fill="x", pady=(0, 10))
        
        ttk.Label(top, text="Έτος:").grid(row=0, column=0, sticky="w")
        ttk.Entry(top, textvariable=self.year_var, width=10).grid(row=0, column=1, sticky="w", padx=(6, 12))
        
        ttk.Label(top, text="Μήνας (1-12):").grid(row=0, column=2, sticky="w")
        ttk.Entry(top, textvariable=self.month_var, width=10).grid(row=0, column=3, sticky="w", padx=(6, 12))
        
        ttk.Label(top, text="Extra αργίες (πχ 25 ή 6,25 ή 6-10):").grid(row=0, column=4, sticky="w")
        ttk.Entry(top, textvariable=self.extra_var, width=18).grid(row=0, column=5, sticky="w", padx=(6, 12))
        
        # Unified "Generate All" button
        ttk.Button(top, text="🌐 Δημιουργία ΟΛΩΝ", command=self.on_generate_all_unified).grid(
            row=0, column=6, sticky="w", padx=(12, 6)
        )
        
        # Export All button
        ttk.Button(top, text="📄 Export ΟΛΩΝ", command=self.on_export_all_word).grid(
            row=0, column=7, sticky="w", padx=(0, 12)
        )
        
        solver_status = "🎯 Backtracking Solver"
        ttk.Label(top, text=solver_status, foreground="blue").grid(
            row=0, column=8, sticky="w", padx=(12, 0)
        )
        
        self.nb = ttk.Notebook(root)
        self.nb.pack(fill="both", expand=True)
        
        self.tabs: dict[str, PeopleTab] = {}
        for key in TAB_KEYS:
            tab = PeopleTab(self.nb, tab_key=key, app_ref=self)
            self.nb.add(tab, text=TAB_TITLES.get(key, key))
            self.tabs[key] = tab
    
    def current_tab_key(self) -> str:
        idx = self.nb.index(self.nb.select())
        return TAB_KEYS[idx]
    
    def get_shared_settings(self):
        year = safe_int(self.year_var.get(), 2026)
        month = safe_int(self.month_var.get(), 3)
        if month < 1 or month > 12:
            raise ValueError("Ο μήνας πρέπει να είναι 1-12.")
        _, dim = calendar.monthrange(year, month)
        extra = parse_days_list(self.extra_var.get(), dim)
        return year, month, extra
    
    def on_generate_all_unified(self):
        """Generate all tabs together (unified scheduling to avoid conflicts)."""
        year, month, extra = self.get_shared_settings()
        _, days_in_month = calendar.monthrange(year, month)
        
        # Collect data from all tabs
        all_tabs_data = {}
        for key in TAB_KEYS:
            try:
                # Call parse_people directly with days_in_month
                names, leaves, ranks, max_caps, preferences = self.tabs[key].parse_people(days_in_month)
                if names:  # Only include non-empty tabs
                    all_tabs_data[key] = {
                        "names": names,
                        "leaves": leaves,
                        "ranks": ranks,
                        "max_caps": max_caps,
                        "preferences": preferences,
                    }
            except Exception as e:
                messagebox.showerror("Σφάλμα", f"Σφάλμα στο {TAB_TITLES[key]}: {str(e)}")
                return
        
        if not all_tabs_data:
            messagebox.showwarning("Προσοχή", "Δεν υπάρχουν δεδομένα σε κανένα tab!")
            return
        
        # Clear all tabs first and start a fresh debug session per tab
        for key, tab in self.tabs.items():
            tab.clear_program_view()
            tab.log("----- Νέα εκτέλεση (ΔΗΜΙΟΥΡΓΙΑ ΟΛΩΝ) -----")
        
        # Run unified scheduling (route debug logs into each tab's right panel)
        try:
            start_key = self.current_tab_key()

            def gui_log(tab_key, message: str):
                target = tab_key or start_key
                if target not in self.tabs:
                    target = start_key
                kind = "error" if str(message).lstrip().startswith("❌") else "info"
                self.tabs[target].log(message, kind=kind)

            results = solve_all_tabs_unified(
                all_tabs_data=all_tabs_data,
                year=year,
                month=month,
                extra_holidays=extra,
                log_cb=gui_log,
            )
            
            # Display results in each tab
            for key, (schedule, quotas, meta, solve_info) in results.items():
                self.tabs[key].render_program(year, month, extra, schedule, 
                                              all_tabs_data[key]["ranks"], quotas, meta, solve_info)
            
            messagebox.showinfo("Επιτυχία", 
                              f"Δημιουργήθηκαν {len(results)} προγράμματα χωρίς συγκρούσεις!")
            
        except ScheduleError as e:
            messagebox.showerror("Αδύνατο πρόγραμμα", f"{e.message}\n\n{e.details}")
        except Exception as e:
            messagebox.showerror("Σφάλμα", f"Απρόσμενο σφάλμα:\n{str(e)}")
    
    def on_export_all_word(self):
        """Export all tabs to a single Word file (4 pages)."""
        year, month, extra = self.get_shared_settings()
        
        # Collect data from all tabs that have schedules
        export_data = []
        for key in TAB_KEYS:
            if self.tabs[key].current_schedule:
                export_data.append({
                    "tab_key": key,
                    "schedule": self.tabs[key].current_schedule,
                    "ranks": self.tabs[key].current_ranks,
                })
        
        if not export_data:
            messagebox.showwarning("Προσοχή", "Δεν υπάρχουν προγράμματα για export!")
            return
        
        # Ask for filename
        default_name = f"Προγραμματα_{GREEK_MONTHS_GEN[month]}_{year}.docx"
        filepath = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Documents", "*.docx")],
            initialfile=default_name
        )
        
        if not filepath:
            return
        
        try:
            export_all_schedules_to_word(
                filepath=filepath,
                export_data=export_data,
                year=year,
                month=month,
                extra_holidays=extra,
            )
            
            # AUTO-SAVE TO HISTORY after successful export
            for data in export_data:
                try:
                    add_to_history(
                        year=year,
                        month=month,
                        tab_key=data["tab_key"],
                        schedule=data["schedule"],
                        ranks=data["ranks"],  # Pass ranks
                        extra_holidays=extra,
                    )
                except Exception as e:
                    print(f"Warning: Could not save {data['tab_key']} to history: {e}")
            
            self.save_schedule_state()
            messagebox.showinfo(
                "Επιτυχία", 
                f"Εξαγωγή σε:\n{filepath}\n\n✅ Αποθηκεύτηκε στο ιστορικό!"
            )
        except Exception as e:
            messagebox.showerror("Σφάλμα", f"Αποτυχία εξαγωγής:\n{str(e)}")
    
    def on_generate_current(self):
        key = self.current_tab_key()
        tab = self.tabs[key]
        try:
            year, month, extra = self.get_shared_settings()
            _, dim = calendar.monthrange(year, month)
            
            tab.clear_program_view()
            tab.log("----- Νέα εκτέλεση -----")
            
            names, leaves, ranks, max_caps, preferences = tab.parse_people(dim)
            
            sched, quotas, meta, solve_info = solve_schedule_best_effort(
                names=names,
                year=year,
                month=month,
                extra_holidays=extra,
                leaves=leaves,
                preferences=preferences,  # NEW: pass preferences
                max_caps=max_caps,
                ranks=ranks,  # NEW: pass ranks for group balancing
                tab_key=key,
                log_cb=lambda m: tab.log(m, kind="info"),
            )
            
            tab.render_program(year, month, extra, sched, ranks, quotas, meta, solve_info)
            messagebox.showinfo("OK", "Βγήκε νέο πρόγραμμα ✅")
        
        except ScheduleError as e:
            tab.log(e.details, kind="error")
            messagebox.showerror("Σφάλμα", e.details)
        except Exception as e:
            import traceback
            full = traceback.format_exc()
            tab.log(f"⚠️ ΕΣΩΤΕΡΙΚΟ ΣΦΑΛΜΑ:\n{full}", kind="error")
            messagebox.showerror("Σφάλμα", f"Εσωτερικό σφάλμα:\n{str(e)}\n\nΔείτε το log για λεπτομέρειες.")

    def on_save_current_schedule(self):
        """Save current tab's schedule to history and state."""
        try:
            year, month, extra = self.get_shared_settings()
            key = self.current_tab_key()
            tab = self.tabs[key]
            if not tab.current_schedule:
                raise RuntimeError("Δεν υπάρχει πρόγραμμα.")
            
            add_to_history(
                year=year,
                month=month,
                tab_key=key,
                schedule=tab.current_schedule,
                ranks=tab.current_ranks,
                extra_holidays=extra,
            )
            self.save_state()
            self.save_schedule_state()
            messagebox.showinfo("Αποθήκευση", f"✅ Αποθηκεύτηκε το πρόγραμμα {TAB_TITLES.get(key, key)}!")
        except Exception as e:
            messagebox.showerror("Σφάλμα", str(e))


    def on_manual_save_state(self):
        """Manually save the current application state to JSON."""
        self.save_state()
        self.save_schedule_state()
        messagebox.showinfo("Αποθήκευση", "✅ Το state αποθηκεύτηκε!")
        
    def save_state(self):
        """Αποθηκεύει ΜΟΝΟ τα inputs (άτομα, ρυθμίσεις) - καλείται στο κλείσιμο και αυτόματα."""
        try:
            people_state = {
                "year": safe_int(self.year_var.get(), 2026),
                "month": safe_int(self.month_var.get(), 3),
                "extra": self.extra_var.get().strip(),
                "active_tab": self.current_tab_key(),
                "tabs": {},
            }
            for key in TAB_KEYS:
                tab_state = self.tabs[key].get_tab_state()
                # Αποθήκευσε ΜΟΝΟ τα people inputs, όχι το schedule
                people_state["tabs"][key] = {
                    "count": tab_state.get("count", 1),
                    "people": tab_state.get("people", []),
                }
            
            with open(get_state_path(), "w", encoding="utf-8") as f:
                json.dump(people_state, f, ensure_ascii=False, indent=2)
        except Exception:
            return

    def save_schedule_state(self):
        """Αποθηκεύει τα παραγόμενα προγράμματα - καλείται ΜΟΝΟ χειροκίνητα."""
        try:
            schedule_state = {"tabs": {}}
            for key in TAB_KEYS:
                tab_state = self.tabs[key].get_tab_state()
                # Αποθήκευσε ΜΟΝΟ τα schedule data
                sched_data = {}
                if "saved_schedule" in tab_state:
                    sched_data["saved_schedule"] = tab_state["saved_schedule"]
                if "saved_ranks" in tab_state:
                    sched_data["saved_ranks"] = tab_state["saved_ranks"]
                if "saved_stats_text" in tab_state:
                    sched_data["saved_stats_text"] = tab_state["saved_stats_text"]
                if sched_data:
                    schedule_state["tabs"][key] = sched_data
            
            with open(get_schedule_path(), "w", encoding="utf-8") as f:
                json.dump(schedule_state, f, ensure_ascii=False, indent=2)
        except Exception:
            return
    
    def load_state(self):
        """Φορτώνει inputs από scheduler_people.json και προγράμματα από scheduler_schedule.json."""
        # --- Φόρτωση inputs (people) ---
        people_path = get_state_path()
        if os.path.exists(people_path):
            try:
                with open(people_path, "r", encoding="utf-8") as f:
                    state = json.load(f)
                
                if "year" in state:
                    self.year_var.set(str(state["year"]))
                if "month" in state:
                    self.month_var.set(str(state["month"]))
                if "extra" in state:
                    self.extra_var.set(state["extra"])
                
                if "tabs" in state:
                    for key in TAB_KEYS:
                        if key in state["tabs"]:
                            self.tabs[key].set_tab_state(state["tabs"][key])
                
                if "active_tab" in state and state["active_tab"] in TAB_KEYS:
                    self.nb.select(TAB_KEYS.index(state["active_tab"]))
            except Exception:
                pass
        
        # --- Φόρτωση προγραμμάτων (schedule) ---
        schedule_path = get_schedule_path()
        if os.path.exists(schedule_path):
            try:
                with open(schedule_path, "r", encoding="utf-8") as f:
                    sched_state = json.load(f)
                
                if "tabs" in sched_state:
                    for key in TAB_KEYS:
                        if key in sched_state["tabs"]:
                            # Φόρτωσε μόνο το schedule, χωρίς να πειράξεις τα people inputs
                            sched_data = sched_state["tabs"][key]
                            tab = self.tabs[key]
                            if "saved_schedule" in sched_data and sched_data["saved_schedule"]:
                                tab.current_schedule = {int(k): v for k, v in sched_data["saved_schedule"].items()}
                                tab.current_ranks = sched_data.get("saved_ranks", {})
                                saved_text = sched_data.get("saved_stats_text", "")
                                if saved_text:
                                    try:
                                        tab.stats.config(state="normal")
                                        tab.stats.delete("1.0", "end")
                                        tab.stats.insert("end", saved_text)
                                        tab.stats.config(state="disabled")
                                        tab.export_btn.config(state="normal")
                                    except Exception:
                                        pass
            except Exception:
                pass
    
    # History Management Methods
    def on_edit_history(self):
        """Open comprehensive history viewer/editor dialog."""
        HistoryViewerEditorDialog(self)
    
    def on_reload_history(self):
        """Reload history from JSON file."""
        try:
            history = load_history()
            messagebox.showinfo("Επιτυχία", "Το ιστορικό φορτώθηκε από το αρχείο!")
        except Exception as e:
            messagebox.showerror("Σφάλμα", f"Αποτυχία φόρτωσης:\n{str(e)}")
    
    def on_open_history_file(self):
        """Open history JSON file in default editor."""
        import subprocess
        import platform
        
        if not os.path.exists(HISTORY_FILE):
            # Create empty history file
            save_history({})
        
        try:
            if platform.system() == 'Darwin':  # macOS
                subprocess.call(('open', HISTORY_FILE))
            elif platform.system() == 'Windows':
                os.startfile(HISTORY_FILE)
            else:  # Linux
                subprocess.call(('xdg-open', HISTORY_FILE))
        except Exception as e:
            messagebox.showerror("Σφάλμα", f"Δεν μπόρεσε να ανοίξει το αρχείο:\n{str(e)}")
    
    def on_clear_history(self):
        """Clear all history after confirmation."""
        if messagebox.askyesno("Επιβεβαίωση", "Σίγουρα θέλεις να διαγράψεις ΟΛΟ το ιστορικό;"):
            save_history({})
            messagebox.showinfo("Επιτυχία", "Το ιστορικό διαγράφηκε!")
    
    def on_close(self):
        self.save_state()
        self.destroy()


# -----------------------------
# History Viewer + Editor Dialog (TABBED BY MONTH)
# -----------------------------
class HistoryViewerEditorDialog(tk.Toplevel):
    def __init__(self, parent):
        super().__init__(parent)
        self.title("Ιστορικό - Προβολή & Επεξεργασία")
        self.geometry("900x750")
        
        # Top: Year selection
        top_frame = ttk.Frame(self, padding=10)
        top_frame.pack(fill="x")
        
        ttk.Label(top_frame, text="Έτος:").pack(side="left", padx=5)
        self.year_var = tk.StringVar(value="2026")
        ttk.Entry(top_frame, textvariable=self.year_var, width=8).pack(side="left", padx=5)
        ttk.Button(top_frame, text="Φόρτωση Έτους", command=self.load_year).pack(side="left", padx=5)
        
        # Create Notebook for tabs
        self.notebook = ttk.Notebook(self)
        self.notebook.pack(fill="both", expand=True, padx=10, pady=5)
        
        # Month tabs (13 tabs: ΤΟΤΑΛ + 12 μήνες)
        self.month_frames = {}
        month_names = ["ΤΟΤΑΛ ΕΤΟΥΣ"] + [GREEK_MONTHS_GEN[i].capitalize() for i in range(1, 13)]
        
        for i, month_name in enumerate(month_names):
            month_key = 0 if i == 0 else i
            
            # Create frame for this month
            month_frame = ttk.Frame(self.notebook)
            self.notebook.add(month_frame, text=month_name)
            
            # Stats label
            stats_label = ttk.Label(month_frame, text="", foreground="blue", font=('', 10, 'bold'))
            stats_label.pack(fill="x", padx=10, pady=5)
            
            # Container for table
            table_container = ttk.Frame(month_frame)
            table_container.pack(fill="both", expand=True, padx=10, pady=5)
            
            # Headers (με monospace font για τέλεια ευθυγράμμιση)
            headers_frame = ttk.Frame(table_container)
            headers_frame.pack(fill="x")
            
            header_font = ('Courier', 9, 'bold')  # Monospace font!
            ttk.Label(headers_frame, text="Όνομα", width=20, font=header_font, anchor="w").grid(row=0, column=0, padx=2)
            ttk.Label(headers_frame, text="Καθημερινές", width=15, font=header_font, anchor="center").grid(row=0, column=1, padx=2)
            ttk.Label(headers_frame, text="Παρασκευές", width=15, font=header_font, anchor="center").grid(row=0, column=2, padx=2)
            ttk.Label(headers_frame, text="Αργίες", width=15, font=header_font, anchor="center").grid(row=0, column=3, padx=2)
            ttk.Label(headers_frame, text="TOTAL", width=10, font=header_font, foreground="darkgreen", anchor="center").grid(row=0, column=4, padx=2)
            
            # Scrollable table
            canvas = tk.Canvas(table_container, height=500)
            scrollbar = ttk.Scrollbar(table_container, orient="vertical", command=canvas.yview)
            scrollable_frame = ttk.Frame(canvas)
            
            scrollable_frame.bind(
                "<Configure>",
                lambda e, c=canvas: c.configure(scrollregion=c.bbox("all"))
            )
            
            canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
            canvas.configure(yscrollcommand=scrollbar.set)
            
            canvas.pack(side="left", fill="both", expand=True)
            scrollbar.pack(side="right", fill="y")
            
            # Store references
            self.month_frames[month_key] = {
                "frame": month_frame,
                "stats_label": stats_label,
                "scrollable_frame": scrollable_frame,
                "rows": []
            }
        
        # Bottom: Buttons
        btn_frame = ttk.Frame(self, padding=10)
        btn_frame.pack(fill="x")
        
        ttk.Button(btn_frame, text="💾 Αποθήκευση Όλων", command=self.save_all_data).pack(side="right", padx=5)
        ttk.Button(btn_frame, text="Κλείσιμο", command=self.destroy).pack(side="right", padx=5)
        
        self.load_year()
    
    def load_year(self):
        """Load all months for the selected year."""
        year = self.year_var.get()
        history = load_history()
        
        if year not in history:
            # Empty year
            for month_key, month_data in self.month_frames.items():
                self.clear_month(month_key)
                if month_key == 0:
                    month_data["stats_label"].config(text=f"📊 ΤΟΤΑΛ ΕΤΟΥΣ {year}: Κενό")
                else:
                    month_name = GREEK_MONTHS_GEN[month_key].capitalize()
                    month_data["stats_label"].config(text=f"📅 {month_name} {year}: Κενό")
            return
        
        year_data = history[year]
        
        # Load ΤΟΤΑΛ (month_key=0)
        self.load_total_year(year_data, year)
        
        # Load each month (month_key=1-12)
        for month_key in range(1, 13):
            self.load_single_month(year_data, year, month_key)
    
    def clear_month(self, month_key):
        """Clear all rows from a month."""
        month_data = self.month_frames[month_key]
        for row_frame, _, _, _, _ in month_data["rows"]:
            row_frame.destroy()
        month_data["rows"] = []
    
    def load_total_year(self, year_data, year):
        """Load TOTAL year view (month_key=0)."""
        month_data = self.month_frames[0]
        self.clear_month(0)
        
        # Aggregate all months
        totals = {}
        for month_str in year_data.keys():
            if not month_str.isdigit():
                continue
            
            for person, data in year_data[month_str].items():
                if person.startswith("_"):  # παράλειψε internal keys όπως _saved_tabs
                    continue
                if person not in totals:
                    totals[person] = {
                        "weekday": 0,
                        "friday": 0,
                        "holiday": 0,
                        "rank": data.get("rank", "")  # Store rank from first occurrence
                    }
                
                totals[person]["weekday"] += data.get("weekday", 0)
                totals[person]["friday"] += data.get("friday", 0)
                totals[person]["holiday"] += data.get("holiday", 0)
        
        # Group by rank
        officers = []
        others = []
        
        for person, data in totals.items():
            rank = data.get("rank", "")
            # Check if rank contains officer title
            if any(suffix in rank.upper() for suffix in ["ΛΓΟΣ", "ΥΠΛΓΟΣ", "ΑΝΘΛΓΟΣ"]):
                officers.append((person, data))
            else:
                others.append((person, data))
        
        month_data["stats_label"].config(
            text=f"📊 ΤΟΤΑΛ ΕΤΟΥΣ {year}: {len(totals)} άτομα (Αξιωματικοί: {len(officers)}, Λοιποί: {len(others)})"
        )
        
        # Display Officers
        if officers:
            sep_frame = ttk.Frame(month_data["scrollable_frame"])
            sep_frame.pack(fill="x", pady=5)
            ttk.Label(sep_frame, text="▼ ΑΞΙΩΜΑΤΙΚΟΙ (ΛΓΟΣ, ΥΠΛΓΟΣ, ΑΝΘΛΓΟΣ) ▼", 
                     font=('', 10, 'bold'), foreground="darkblue").pack()
            
            for person, data in sorted(officers):
                self.add_row_to_month(0, person, data["weekday"], data["friday"], data["holiday"], editable=False)
        
        # Display Others
        if others:
            sep_frame = ttk.Frame(month_data["scrollable_frame"])
            sep_frame.pack(fill="x", pady=5)
            ttk.Label(sep_frame, text="▼ ΛΟΙΠΟΙ ▼", 
                     font=('', 10, 'bold'), foreground="darkgreen").pack()
            
            for person, data in sorted(others):
                self.add_row_to_month(0, person, data["weekday"], data["friday"], data["holiday"], editable=False)
    
    def load_single_month(self, year_data, year, month_key):
        """Load single month view."""
        month_data = self.month_frames[month_key]
        self.clear_month(month_key)
        
        month_str = str(month_key)
        month_name = GREEK_MONTHS_GEN[month_key].capitalize()
        
        if month_str not in year_data:
            month_data["stats_label"].config(text=f"📅 {month_name} {year}: Κενό")
            return
        
        people_data = year_data[month_str]
        
        # Group people by rank
        officers = []
        others = []
        
        for person, data in people_data.items():
            if person.startswith("_"):  # παράλειψε internal keys όπως _saved_tabs
                continue
            rank = data.get("rank", "")
            # Check if rank contains officer title
            if any(suffix in rank.upper() for suffix in ["ΛΓΟΣ", "ΥΠΛΓΟΣ", "ΑΝΘΛΓΟΣ"]):
                officers.append((person, data))
            else:
                others.append((person, data))
        
        month_data["stats_label"].config(
            text=f"📅 {month_name} {year}: {len(officers)+len(others)} άτομα (Αξ: {len(officers)}, Λοιποί: {len(others)})"
        )
        
        # Display Officers
        if officers:
            sep_frame = ttk.Frame(month_data["scrollable_frame"])
            sep_frame.pack(fill="x", pady=5)
            ttk.Label(sep_frame, text="▼ ΑΞΙΩΜΑΤΙΚΟΙ ▼", 
                     font=('', 9, 'bold'), foreground="darkblue").pack()
            
            for person, data in sorted(officers):
                self.add_row_to_month(month_key, person, data.get("weekday", 0), data.get("friday", 0), data.get("holiday", 0), editable=True)
        
        # Display Others
        if others:
            sep_frame = ttk.Frame(month_data["scrollable_frame"])
            sep_frame.pack(fill="x", pady=5)
            ttk.Label(sep_frame, text="▼ ΛΟΙΠΟΙ ▼", 
                     font=('', 9, 'bold'), foreground="darkgreen").pack()
            
            for person, data in sorted(others):
                self.add_row_to_month(month_key, person, data.get("weekday", 0), data.get("friday", 0), data.get("holiday", 0), editable=True)
    
    def add_row_to_month(self, month_key, name="", weekday=0, friday=0, holiday=0, editable=True):
        """Add a row to specific month."""
        month_data = self.month_frames[month_key]
        
        row_frame = ttk.Frame(month_data["scrollable_frame"])
        row_frame.pack(fill="x", pady=1)
        
        name_var = tk.StringVar(value=name)
        weekday_var = tk.StringVar(value=str(weekday))
        friday_var = tk.StringVar(value=str(friday))
        holiday_var = tk.StringVar(value=str(holiday))
        
        # Monospace font για τέλεια ευθυγράμμιση!
        data_font = ('Courier', 9)
        
        # Name
        if editable:
            tk.Entry(row_frame, textvariable=name_var, width=20, font=data_font).grid(row=0, column=0, padx=2)
        else:
            ttk.Label(row_frame, text=name, width=20, font=data_font, anchor="w").grid(row=0, column=0, padx=2)
        
        # Weekday (centered με monospace)
        if editable:
            tk.Entry(row_frame, textvariable=weekday_var, width=15, justify="center", font=data_font).grid(row=0, column=1, padx=2)
        else:
            ttk.Label(row_frame, text=str(weekday), width=15, font=data_font, anchor="center").grid(row=0, column=1, padx=2)
        
        # Friday (centered με monospace)
        if editable:
            tk.Entry(row_frame, textvariable=friday_var, width=15, justify="center", font=data_font).grid(row=0, column=2, padx=2)
        else:
            ttk.Label(row_frame, text=str(friday), width=15, font=data_font, anchor="center").grid(row=0, column=2, padx=2)
        
        # Holiday (centered με monospace)
        if editable:
            tk.Entry(row_frame, textvariable=holiday_var, width=15, justify="center", font=data_font).grid(row=0, column=3, padx=2)
        else:
            ttk.Label(row_frame, text=str(holiday), width=15, font=data_font, anchor="center").grid(row=0, column=3, padx=2)
        
        # Total (always read-only, centered με monospace)
        total = weekday + friday + holiday
        ttk.Label(row_frame, text=str(total), width=10, foreground="darkgreen", 
                 font=('Courier', 9, 'bold'), anchor="center").grid(row=0, column=4, padx=2)
        
        if editable:
            month_data["rows"].append((row_frame, name_var, weekday_var, friday_var, holiday_var))
    
    def save_all_data(self):
        """Save all modified months."""
        history = load_history()
        year = self.year_var.get()
        
        if year not in history:
            history[year] = {}
        
        # Save each month (skip TOTAL which is month_key=0)
        for month_key in range(1, 13):
            month_data = self.month_frames[month_key]
            month_str = str(month_key)
            
            if month_str not in history[year]:
                history[year][month_str] = {}
            
            # Clear and rebuild
            history[year][month_str] = {}
            
            for _, name_var, weekday_var, friday_var, holiday_var in month_data["rows"]:
                name = name_var.get().strip()
                if not name:
                    continue
                
                try:
                    weekday = int(weekday_var.get())
                    friday = int(friday_var.get())
                    holiday = int(holiday_var.get())
                    
                    if weekday > 0 or friday > 0 or holiday > 0:
                        history[year][month_str][name] = {
                            "weekday": weekday,
                            "friday": friday,
                            "holiday": holiday,
                        }
                except ValueError:
                    messagebox.showerror("Σφάλμα", f"Μη έγκυρος αριθμός για {name}")
                    return
        
        save_history(history)
        messagebox.showinfo("Επιτυχία", f"Όλα τα δεδομένα αποθηκεύτηκαν!")
    

if __name__ == "__main__":
    app = SchedulerApp()
    app.mainloop()